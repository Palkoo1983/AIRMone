
def pick_year_columns(headers):
    """Universal rule for 3+ grouped columns:
    - If headers include a middle 'Módosítások'/'Helyesbítés'/'Korrekció' column, ignore it.
    - Return (prev_idx, curr_idx) where prev is the first numeric-like col and curr is the last.
    """
    names = [str(h).lower() for h in headers]
    # strip known middle columns
    ignored = {"módosítások","modositasok","helyesbítés","helyesbites","korrekció","korrekcio"}
    clean = [(i,n) for i,n in enumerate(names) if n.strip() not in ignored]
    if len(clean) >= 2:
        prev_idx = clean[0][0]
        curr_idx = clean[-1][0]
        return prev_idx, curr_idx
    # fallback: first and last
    return 0, max(0, len(headers)-1)


# --- Robust number extraction for space-grouped thousands (handles negatives and parentheses) ---
def _extract_grouped_numbers(line: str):
    """
    Extract a list of integers from a line.
    Handles both plain numbers (e.g., 3815087) and thousands‑grouped with spaces/dots/nbsp
    (e.g., "510 432", "155 474", "1.234.567").
    Crucially: it DOES NOT merge two adjacent amounts like "510 432 155 474".
    """
    if not line:
        return []

    import re
    s = line.replace('\u00A0', ' ')
    # Pattern: either grouped by thousands with separator, or a plain integer
    pat = re.compile(r'[-+−]?(?:\d{1,3}(?:[ .]\d{3})+|\d+)', re.UNICODE)
    nums = []
    for m in pat.finditer(s):
        token = m.group(0).replace('−', '-').replace(' ', '').replace('.', '')
        try:
            val = int(token)
        except Exception:
            continue
        nums.append(val)
    return nums

def prev_curr_from_line(s: str):
    nums = _extract_grouped_numbers(s)
    if not nums:
        return None, None
    if len(nums) == 1:
        return None, nums[0]
    return nums[0], nums[-1]

# --- Universal extractor for "previous / (modifications) / current" rows ---
def _nums_from_text(s: str):
    if not s: return []
    s = s.replace('\u2212','-')  # unicode minus
    out = []
    for m in re.finditer(r'\(?-?\d[\d\s]*\)?', s):
        token = m.group(0).strip()
        if not re.search(r'\d', token):
            continue
        neg = token.startswith('(') and token.endswith(')')
        token = token.strip('()').replace(' ','')
        try:
            val = int(token)
            if neg: val = -val
            out.append(val)
        except:
            pass
    return out

def prev_curr_from_line(s: str):
    nums = _nums_from_text(s)
    if not nums:
        return None, None
    if len(nums) == 1:
        return None, nums[0]
    return nums[0], nums[-1]

# --- English label map for Key financials (module-level) ---
label_en = {
    "Értékesítés nettó árbevétele":"Net sales revenue",
    "Anyagjellegű ráfordítások":"Material-type expenses",
    "Személyi jellegű ráfordítások":"Personnel expenses",
    "Értékcsökkenési leírás":"Depreciation and amortization",
    "Egyéb bevételek":"Other income",
    "Egyéb ráfordítások":"Other expenses",
    "Adózott eredmény":"Profit after tax",
    "Üzemi (üzleti) tevékenység eredménye":"Operating profit (loss)",
    "Forgóeszközök":"Current assets",
    "Készletek":"Inventory",
    "Követelések":"Receivables",
    "Pénzeszközök":"Cash and cash equivalents",
    "Szállítók":"Payables",
    "Rövid lejáratú kötelezettségek":"Short-term liabilities",
    "Hosszú lejáratú kötelezettségek":"Long-term liabilities",
    "Kötelezettségek összesen":"Total liabilities",
    "Saját tőke":"Equity"
}


import airm_hotfix_universal  # UNIVERSAL HOTFIX – do not remove
import re, sys, json, unicodedata
from pathlib import Path


# --- Minimal negative handling (safe, localized) ---
MINUS_CHARS = "\u2212\u2012\u2013\u2014"
SPACE_CHARS = (" ", "\xa0", "\u202f")
def parse_int_signed(token: str):
    if token is None:
        return None
    s = str(token).strip()
    neg = False
    if s.startswith("(") and s.endswith(")"):
        neg = True
        s = s[1:-1].strip()
    if s and (s[0] in "+-" or s[0] in MINUS_CHARS):
        if s[0] in "-" or s[0] in MINUS_CHARS:
            neg = not neg
        s = s[1:].strip()
    for ch in SPACE_CHARS:
        s = s.replace(ch, "")
    s = s.replace(".", "")
    if not s.isdigit():
        return None
    v = int(s)
    return -v if neg else v
SIGNED_NUM = r"\(?[+\-\u2212\u2012\u2013\u2014]?\s*(?:\d{1,3}(?:[\s\xa0]\d{3})+|\d+)\)?"
SIGNED_PAIR_AT_END_RE = re.compile(rf"({SIGNED_NUM})[\s\xa0]+({SIGNED_NUM})[\s\xa0]*$")

SIGNED_GROUPED_PAIR_AT_END_RE = re.compile(r"([()\+\-\u2212\u2012\u2013\u2014]?\s*\d{1,3}(?:[\s\xa0]\d{3})+)\s+([()\+\-\u2212\u2012\u2013\u2014]?\s*\d{1,3}(?:[\s\xa0]\d{3})+)\s*$")
# --- /Minimal negative handling ---

def strip_accents(s: str) -> str:
    return ''.join(ch for ch in unicodedata.normalize('NFKD', s) if not unicodedata.combining(ch))

NUM_RE = re.compile(r'(?:\d{1,3}(?:[\s\xa0]\d{3})+|\d+)')
PAIR_AT_END_RE = re.compile(r'((?:\d{1,3}(?:[\s\xa0]\d{3})+?|\d+))[\s\xa0]+((?:\d{1,3}(?:[\s\xa0]\d{3})+|\d+))[\s\xa0]*$')

def read_pdf_text(pdf_path: Path) -> str:
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(str(pdf_path)) as pdf:
            parts = []
            globals()['_AIRM_LAST_WORDS'] = []
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
                try:
                    words = page.extract_words(use_text_flow=False, keep_blank_chars=False)
                except TypeError:
                    words = page.extract_words()
                globals()['_AIRM_LAST_WORDS'].append(words)
            text = "\n".join(parts)
    except Exception:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(pdf_path))
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
            text = "\n".join(parts)
        except Exception as e2:
            raise RuntimeError(f"Nem sikerült beolvasni a PDF-et: {e2}")
    return text











def get_suppliers_from_pdf101(text: str):
    """Return (current, previous, line) for 101. Szállítók. Handles 'prev curr' concatenation and ignores the 101 code."""
    if not text:
        return None, None, None
    import re
    lines = text.splitlines()
    pat = re.compile(r'^\s*101[.)]?\s')
    i101 = None
    for i, line in enumerate(lines):
        if pat.match(line):
            i101 = i
            break
    if i101 is None:
        for i, line in enumerate(lines):
            if re.search(r'\(\s*szállítók\s*\)', line, flags=re.IGNORECASE):
                if i > 0 and pat.match(lines[i-1]):
                    i101 = i-1
                    break
                if i + 1 < len(lines) and pat.match(lines[i+1]):
                    i101 = i+1
                    break
    if i101 is None:
        return None, None, None
    line = lines[i101]

    # Strip leading '101.' or '101 ' so grouping won't include it
    line_wo_code = re.sub(r'^\s*101[.)]?\s*', '', line)

    nums = _extract_grouped_numbers(line_wo_code)
    if not nums:
        return None, None, None

    # If exactly two numbers parsed, take them directly
    clean = [n for n in nums if n is not None and abs(n) >= 1000]
    if len(clean) == 2:
        prev, curr = clean[0], clean[1]
        return curr, prev, line

    # Heuristic: detect concatenation by even count of 3-digit groups on the code-stripped tail
    try:
        grp = re.findall(r'\d{1,3}', line_wo_code)
        if len(grp) >= 4 and len(grp) % 2 == 0:
            half = len(grp)//2
            left = int(''.join(grp[:half]))
            right = int(''.join(grp[half:]))
            if left >= 1000 and right >= 1000:
                return right, left, line
    except Exception:
        pass

    # Fallback: first/last from the filtered list
    if len(clean) >= 2:
        prev, curr = clean[0], clean[-1]
    elif len(clean) == 1:
        prev, curr = None, clean[0]
    else:
        prev, curr = None, None
    return curr, prev, line





def get_suppliers_by_label(text: str):
    """Return (current, previous, line) using the explicit Suppliers label.
    Robust when the label and the numbers are on adjacent lines and when a middle 'Módosítások' column exists.
    """
    if not text:
        return None, None, None
    import re, unicodedata
    def deburr(s: str):
        s = s.replace("\u00A0", " ")
        nf = unicodedata.normalize("NFD", s)
        return "".join(ch for ch in nf if unicodedata.category(ch) != "Mn").lower()
    aliases = ["szallitok", "szallito", "aruszallitasbol", "trade payables", "accounts payable"]
    lines = text.splitlines()
    best_i = None
    for i, ln in enumerate(lines):
        s = deburr(ln)
        if any(a in s for a in aliases):
            if "(szallitok)" in s or "trade payables" in s:
                best_i = i
                break
            if best_i is None:
                best_i = i
    if best_i is None:
        return None, None, None
    # Candidate lines to probe for numbers: label line and its neighbors
    candidates = []
    for j in (best_i, best_i-1, best_i+1):
        if 0 <= j < len(lines):
            candidates.append(lines[j])
    for cand in candidates:
        # Strip leading codes like '111.' before parsing
        s = re.sub(r'^\s*\d+[.)]?\s*', '', cand)
        # Use the general parser: handles concatenated chains and 3+ group cases
        cur, prev = current_year_value_from_line(s)
        if cur is not None or prev is not None:
            # Final guard: if 3+ grouped numbers present, force (prev=first, curr=last)
            grp = re.findall(r'[+\-\u2212\u2012\u2013\u2014]?\s*\d{1,3}(?:[ \xa0]\d{3})+', cand)
            if len(grp) >= 3:
                cur = parse_int_signed(grp[-1]); prev = parse_int_signed(grp[0])
            return cur, prev, cand
    return None, None, Nones[best_i]


def get_trade_payables_dual_universal(text: str):
    """
    Return (current, previous, line) for 'Szállítók' detected by label aliases (not tied to '101.' code).
    - Finds the first line that looks like a Suppliers label (Hungarian/English variants).
    - Extracts up to two large integers (>=1,000). If two found, assumes left=previous, right=current.
    - If only one found, treats it as current and leaves previous=None.
    - Handles concatenated patterns like "510 432 155 474" by splitting into two halves.
    """
    if not text:
        return None, None, None
    import re, unicodedata
    def deburr(s: str):
        s = s.replace("\\u00A0"," ")
        nf = unicodedata.normalize("NFD", s)
        return "".join(ch for ch in nf if unicodedata.category(ch) != "Mn").lower()

    aliases = ["szallito", "szallitok", "aruszallitasbol", "accounts payable", "trade payables"]
    lines = text.splitlines()
    dlines = [deburr(ln) for ln in lines]
    cand_idx = [i for i, dl in enumerate(dlines) if any(a in dl for a in aliases)]
    if not cand_idx:
        return None, None, None
    li = cand_idx[0]
    line = lines[li]
    # try same-line first, then neighbor lines
    window = [x for x in [li, li+1 if li+1 < len(lines) else None, li-1 if li-1 >= 0 else None] if x is not None]

    def extract_two_numbers(s: str):
        nums = _extract_grouped_numbers(s)
        clean = [n for n in nums if n is not None and abs(n) >= 1000]
        if len(clean) >= 2:
            return clean[-1], clean[0]  # (current, previous) assuming rightmost is current
        # concatenation heuristic on stripped non-letters
        try:
            grp = re.findall(r'\\d{1,3}', re.sub(r'\\D+', '', s))
            if len(grp) >= 4 and len(grp) % 2 == 0:
                half = len(grp)//2
                left = int(''.join(grp[:half]))
                right = int(''.join(grp[half:]))
                if left >= 1000 and right >= 1000:
                    return right, left
        except Exception:
            pass
        if len(clean) == 1:
            return clean[0], None
        return None, None

    for idx in window:
        cur, prev = extract_two_numbers(lines[idx])
        if cur is not None or prev is not None:
            return cur, prev, lines[idx]
    return None, None, None
def get_trade_payables_universal(text: str):
    """
    Universal 'Szállítók' (trade payables) extraction based on label aliases; not tied to line numbers.
    v12: +neighbor-line join window; magnitude filter (>=1,000); glued-token split; diacritic-insensitive.
    """
    def deburr(s: str):
        s = s.replace("\u00A0"," ")
        nf = unicodedata.normalize("NFD", s)
        return "".join(ch for ch in nf if unicodedata.category(ch) != "Mn").lower()

    aliases = ["szallito", "szallitok", "aruszallitasbol", "accounts payable", "trade payables"]
    lines = text.splitlines()
    # Precompute deburred lines
    dlines = [deburr(ln) for ln in lines]
    cand_idx = [i for i, dl in enumerate(dlines) if any(a in dl for a in aliases)]
    if not cand_idx:
        return None

    def parse_int(s: str):
        s = s.replace("\u00A0"," ").replace(" ", "").replace(",", "").replace(".", "")
        m = re.findall(r"-?\d+", s)
        if not m:
            return None
        try:
            return int(m[-1])
        except:
            return None

    def split_glued_current(s: str):
        groups = re.findall(r"\d{1,3}", s.replace("\u00A0"," "))
        if len(groups) >= 4:
            return int("".join(groups[-2:]))
        return None

    NUM_GROUPED = re.compile(r"(?:\d{1,3}(?:[ \u00A0]\d{3})+)")
    NUM_ANY     = re.compile(r"\d+")

    def pick_last_reasonable_number(s: str):
        # 1) Grouped tokens first: handle glued prev+current per token
        grouped = NUM_GROUPED.findall(s)
        if grouped:
            vals = []
            for tok in grouped:
                g = split_glued_current(tok)
                if g is not None and abs(g) >= 1000:
                    vals.append(g)
                    continue
                v = parse_int(tok)
                if v is not None and abs(v) >= 1000:
                    vals.append(v)
            if vals:
                return vals[-1]
        # 2) Fallback to any integers (filter tiny/section-like)
        anynums = [int(n) for n in NUM_ANY.findall(s) if n.isdigit()]
        anynums = [n for n in anynums if abs(n) >= 1000]
        if anynums:
            return anynums[-1]
        return None

    # Walk from bottom-most candidate upwards; join with neighbors to catch line-break splits
    for i in reversed(cand_idx):
        window = lines[i]
        if i+1 < len(lines):
            window = window + " " + lines[i+1]
        if i-1 >= 0:
            window = lines[i-1] + " " + window
        v = pick_last_reasonable_number(window)
        if v is not None:
            return v

    # As a last resort, scan the whole text for a '(szállítók' block and apply the same logic
    m = re.search(r"\(\s*sz[aá]ll[ií]t[oó]k\s*\).*?(\d[\d \u00A0]{3,,})", deburr(text), flags=re.DOTALL)
    if m:
        v = pick_last_reasonable_number(m.group(0))
        if v is not None:
            return v
    return None

    def parse_int(s: str):
        s = s.replace("\u00A0"," ").replace(" ", "").replace(",", "").replace(".", "")
        m = re.findall(r"-?\d+", s)
        if not m: 
            return None
        try:
            return int(m[-1])
        except:
            return None

    def split_glued_current(s: str):
        # If a long token contains many 3-digit groups, assume last 2 groups are current-year
        groups = re.findall(r"\d{1,3}", s.replace("\u00A0"," "))
        if len(groups) >= 4:
            cur = int("".join(groups[-2:]))
            return cur
        return None

    NUM_GROUPED = re.compile(r"(?:\d{1,3}(?:[ \u00A0]\d{3})+)")
    NUM_ANY     = re.compile(r"\d+")

    def pick_last_reasonable_number(line: str):
        # 1) Try glued groups
        g = split_glued_current(line)
        if g is not None and abs(g) >= 1000:
            return g
        # 2) Try grouped numbers (e.g., "2 064 948")
        grouped = NUM_GROUPED.findall(line)
        if grouped:
            vals = []
            for tok in grouped:
                v = parse_int(tok)
                if v is not None and abs(v) >= 1000:
                    vals.append(v)
            if vals:
                return vals[-1]
        # 3) Fallback to any integers but drop tiny/section-like tokens
        anynums = NUM_ANY.findall(line)
        anynums = [int(n) for n in anynums if n.isdigit()]
        anynums = [n for n in anynums if abs(n) >= 1000]  # drop 1..999 (e.g., "4" from "111.4")
        if anynums:
            return anynums[-1]
        return None

    # Walk from bottom to top to prefer current-year blocks
    for ln in reversed(cand):
        v = pick_last_reasonable_number(ln)
        if v is not None:
            return v
    return None
    def extract_value(s: str):
        # Find all number groups
        nums = re.findall(r'[0-9][0-9 \u00A0]{2,}[0-9]', s)
        candidates = []
        for raw in nums:
            digits = re.sub(r'\D','', raw)
            if not digits: continue
            try:
                val = int(digits)
                # Filter out unrealistically small (likely line codes)
                if val >= 1000:
                    candidates.append(val)
            except: 
                pass
        if candidates:
            return candidates[-1]
        # fallback: split glued 3-digit groups
        g = re.findall(r'\d{1,3}', s)
        if len(g) >= 4:
            try:
                val = int("".join(g[-2:]))
                if val >= 1000:
                    return val
            except:
                pass
        return None
    for ln in reversed(cand):
        v = extract_value(ln)
        if v is not None:
            return v
    return None
    def split_glued_numbers(s: str):
        g = re.findall(r'\d{1,3}', s)
        if len(g) >= 4:
            # assume last two groups are current-year (e.g., ... 2 064 948 959 928)
            return int("".join(g[-2:]))
        m = re.findall(r'\d+', s)
        return int(m[-1]) if m else None
    for ln in reversed(cand):
        v = split_glued_numbers(ln)
        if v is not None:
            return v
    return None
    NUM = re.compile(r"[\(]?\s*[-−]?\s*(?:\d{1,3}(?:[ \u00A0]\d{3})+|\d+)(?:[.,]\d+)?\s*[\)]?")
    def parse_num(tok: str):
        # Heuristic: some PDFs glue previous & current into one token like '2 064 948 959 928'
        groups = re.findall(r"\d{1,3}", tok)
        if len(groups) >= 4:
            cur_groups = groups[-2:]  # assume last 2 groups are current year
            try:
                return int("".join(cur_groups))
            except Exception:
                pass
        neg = "(" in tok and ")" in tok
        tok = tok.replace("−", "-").replace("(", "").replace(")", "").replace("\u00A0"," ")
        tok = tok.replace(" ", "").replace(",", ".")
        try:
            val = float(tok)
        except:
            digits = re.sub(r"\D","", tok)
            if not digits: return None
            val = int(digits)
        return int(round(-val if neg else val))
    for ln in reversed(cand):
        nums = NUM.findall(ln)
        if nums:
            v = parse_num(nums[-1])
            if v is not None:
                return v
    return None
    NUM = re.compile(r"[\(]?\s*[-−]?\s*(?:\d{1,3}(?:[ \u00A0]\d{3})+|\d+)(?:[.,]\d+)?\s*[\)]?")
    def parse_num(tok: str):
        neg = "(" in tok and ")" in tok
        tok = tok.replace("−", "-").replace("(", "").replace(")", "").replace("\u00A0"," ")
        tok = tok.replace(" ", "").replace(",", ".")
        try:
            val = float(tok)
        except:
            digits = re.sub(r"\D","", tok)
            if not digits: return None
            val = int(digits)
        return int(round(-val if neg else val))
    for ln in reversed(cand):
        nums = NUM.findall(ln)
        if nums:
            v = parse_num(nums[-1])
            if v is not None:
                return v
    return None




def segment_sections(text: str):
    bal_start = re.search(r'MÉRLEGE', text, re.IGNORECASE)
    pl_start = re.search(r'EREDMÉNYKIMUTATÁS', text, re.IGNORECASE)
    if not (bal_start and pl_start):
        return text, text
    if bal_start.start() < pl_start.start():
        return text[bal_start.start():pl_start.start()], text[pl_start.start():]
    return text, text

def numbers_on_line(line: str):
    line = line.replace("\xa0", " ")
    line = re.sub(r'^\s*\d+\.\s*', '', line)
    return NUM_RE.findall(line)




def current_year_value_from_line(line: str):
    line = line.replace("\xa0", " ")
    # 0) UNIVERSAL grouped-number rule (ignore small codes like '020')
    _grp = re.findall(r'[+\-\u2212\u2012\u2013\u2014]?\s*\d{1,3}(?:[ \xa0]\d{3})+', line)
    if len(_grp) >= 3:
        return parse_int_signed(_grp[-1]), parse_int_signed(_grp[0])
    if len(_grp) == 1:
        # single long chain like '2 064 948 959 928' -> try split into two numbers
        parts = re.findall(r'\d{1,3}', _grp[0])
        for cut in range(len(parts)-2, 1, -1):
            left = ' '.join(parts[:cut]); right = ' '.join(parts[cut:])
            lv = parse_int_signed(left); rv = parse_int_signed(right)
            if lv is not None and rv is not None and abs(lv) >= 1000 and abs(rv) >= 1000:
                return rv, lv
    if len(_grp) == 2:
        # Case A: FIRST token contains prev+mod concatenated, second token is current (e.g., '510 432 155 474' & '-84 928')
        g0 = re.findall(r'\d{1,3}', _grp[0])
        if len(g0) >= 4:
            for pg in (3, 2):  # prefer 3-group prev, else 2-group
                if len(g0) - pg in (2, 3) and pg <= len(g0)-2:
                    prev = parse_int_signed(' '.join(g0[:pg]))
                    cur = parse_int_signed(_grp[1])
                    return cur, prev
        # Case B: SECOND token contains mod+current glued; split half
        g2 = re.findall(r'\d{1,3}', _grp[1])
        if len(g2) >= 4 and len(g2) % 2 == 0:
            cur = parse_int_signed(' '.join(g2[len(g2)//2:]))
            prev = parse_int_signed(_grp[0])
            return cur, prev
        return parse_int_signed(_grp[1]), parse_int_signed(_grp[0])
    # 1) Two grouped numbers at end (each has at least one thousand separator)
    m0 = SIGNED_GROUPED_PAIR_AT_END_RE.search(line)
    if m0:
        return parse_int_signed(m0.group(2)), parse_int_signed(m0.group(1))
    # 2) Generic signed pair at end
    m = SIGNED_PAIR_AT_END_RE.search(line)
    if m:
        return parse_int_signed(m.group(2)), parse_int_signed(m.group(1))
    # 3) Tail group pairing (fallback for uneven groups like '2 972 773 995 413')
    tail_m = re.search(r'([()\+\-\u2212\u2012\u2013\u2014\s\d]+)$', line)
    if tail_m:
        tail = tail_m.group(1)
        groups = re.findall(r'\d{1,3}', tail)
        if len(groups) >= 4:
            cur_g = groups[-2:]
            prev_g = groups[:-2]
            def _merge(gs):
                if not gs: return None
                s = gs[0]
                for g in gs[1:]:
                    s += g.rjust(3, "0")
                return parse_int_signed(s)
            cur = _merge(cur_g)
            prev = _merge(prev_g)
            if cur is not None and prev is not None:
                # infer minus sign for current if there is a trailing dash before last number
                if re.search(r'[\-\u2212\u2012\u2013\u2014]\s*\d{1,3}\s*\d{3}\s*$', tail):
                    cur = -abs(cur)
                if "(" in tail and tail.rfind("(") > tail.rfind(")"):
                    cur = -abs(cur)
                return cur, prev
    # 4) Tokens fallback
    nums = numbers_on_line(line)
    if not nums:
        return None, None
    cur = parse_int_signed(nums[-1])
    prev = parse_int_signed(nums[-2]) if len(nums) >= 2 else None
    return cur, prev

def find_line(section_text: str, label_regex: str):
    # Prefer matches with numbers (ideally 2+ grouped numbers). Skip text-only hits like 'felhasználás'.
    pat = re.compile(label_regex, re.IGNORECASE)
    best = None
    best_score = -1
    for line in section_text.splitlines():
        if pat.search(line):
            g = re.findall(r'[+\-\u2212\u2012\u2013\u2014]?\s*\d{1,3}(?:[ \xa0]\d{3})+', line)
            n = re.findall(r'\d+', line)
            score = len(g)*2 + (1 if n else 0)
            if score > 0:
                cur, prev = current_year_value_from_line(line)
                if len(g) >= 3:
                    cur = parse_int_signed(g[-1])
                    prev = parse_int_signed(g[0])
                return {"line": line, "current": cur, "previous": prev}
            if score > best_score:
                best = line; best_score = score
    if best is not None:
        cur, prev = current_year_value_from_line(best)
        return {"line": best, "current": cur, "previous": prev}
    return {"line": None, "current": None, "previous": None}


def find_revenue_line(section_text: str):
    for line in section_text.splitlines():
        s = strip_accents(line).lower()
        if 'ertekesites' in s and 'netto' in s and 'arbev' in s and 'belfoldi' not in s and 'export' not in s and (' i.' in s or s.strip().startswith('i.')):
            cur, prev = current_year_value_from_line(line)
            if cur is not None:
                return {"line": line, "current": cur, "previous": prev}
    return {"line": None, "current": None, "previous": None}

KEYS_BS = [
    ("Forgóeszközök", r'\bB\.\s*Forgóeszközök\b|\bForgóeszközök\b'),
    ("Készletek", r'^\s*\d+\.\s*I\.\s*Készletek\b|\bKészletek\b'),
    ("Követelések", r'^\s*\d+\.\s*II\.\s*Követelések\b|\bKövetelések\b'),
    ("Pénzeszközök", r'^\s*\d+\.\s*IV\.\s*Pénzeszközök\b|\bPénzeszközök\b'),
    ("Eszközök összesen", r'Eszközök\s*\(aktívák\)\s*összesen|Eszközök.*összesen'),
    ("Saját tőke", r'^\s*\d+\.\s*D\.\s*Saját tőke\b|\bSaját tőke\b'),
    ("Hosszú lejáratú kötelezettségek", r'^\s*\d+\.\s*II\.\s*Hosszú lejáratú kötelezettségek\b|\bHosszú lejáratú kötelezettségek\b'),
    ("Rövid lejáratú kötelezettségek", r'^\s*\d+\.\s*III\.\s*Rövid lejáratú kötelezettségek\b|\bRövid lejáratú kötelezettségek\b'),
    ("Kötelezettségek összesen", r'^\s*\d+\.\s*F\.\s*Kötelezettségek\b'),
]

KEYS_PL = [
    ("Üzemi (üzleti) tevékenység eredménye", r'Üzemi\s*\(üzleti\)\s*tevékenység\s*eredménye'),
    ("Értékesítés nettó árbevétele", r'(?:(?:^\s*\d+\.\s*)?I\.\s*)?Értékesítés[\s\-]+nett[oő][\s\-]*árbevétel(?:e|e)?\b|Értékesítés.*?nett[oő].{0,3}árbev\w+'),
    ("Anyagjellegű ráfordítások", r'^\s*\d+\.\s*IV\.\s*Anyagjellegű ráfordítások|Anyagjellegű ráfordítások'),
    ("Személyi jellegű ráfordítások", r'^\s*\d+\.\s*V\.\s*Személyi jellegű ráfordítások|Személyi jellegű ráfordítások'),
    ("Értékcsökkenési leírás", r'Értékcsökkenési leírás'),
    ("Egyéb bevételek", r'Egyéb bevételek'),
    ("Egyéb ráfordítások", r'Egyéb ráfordítások'),
    ("Adózott eredmény", r'^\s*\d+\.\s*D\.\s*Adózott eredmény|\bAdózott eredmény\b'),
]

def parse_financials_with_raw(text: str):
    bal, pl = segment_sections(text)
    raw = {"balance": {}, "pl": {}}
    bs = {}
    for key, rgx in KEYS_BS:
        info = find_line(bal, rgx)
        raw["balance"][key] = info
        bs[key] = info["current"]
    plv = {}
    for key, rgx in KEYS_PL:
        info = find_line(pl, rgx)
        if key == "Értékesítés nettó árbevétele" and (info["current"] is None):
            info = find_revenue_line(pl)
        raw["pl"][key] = info
        plv[key] = info["current"]
    
    # --- Inject: ensure Szállítók (101. sor) is parsed from the BALANCE by 101 line ---
    # FIRST: label-based suppliers (robust to 111. row and variations)
    try:
        cur_sup2, prev_sup2, sup_line2 = get_suppliers_by_label(text)
        if (cur_sup2 is not None) or (prev_sup2 is not None):
            raw['balance']['Szállítók'] = {'line': sup_line2, 'current': cur_sup2, 'previous': prev_sup2}
            bs['Szállítók'] = cur_sup2 if isinstance(cur_sup2,(int,float)) else (int(str(cur_sup2).replace(' ', '')) if cur_sup2 not in (None,'') else None)
    except Exception:
        pass

    # then try the legacy 101-based detector
    try:
        cur_sup, prev_sup, sup_line = get_suppliers_from_pdf101(text)
        if (cur_sup is not None) or (prev_sup is not None):
            raw["balance"]["Szállítók"] = {"line": sup_line, "current": cur_sup, "previous": prev_sup}
            bs["Szállítók"] = cur_sup if isinstance(cur_sup,(int,float)) else (int(str(cur_sup).replace(" ", "")) if cur_sup not in (None,"") else None)
    except Exception:
        # leave as-is if not found
        pass
    # --- AUTO 'Szállítók' universal + 101 fallback (non-invasive) ---
    try:
        auto_sup = get_trade_payables_universal(text)
        if auto_sup is None:
            auto_sup = get_suppliers_from_pdf101(text)
    except Exception:
        auto_sup = get_suppliers_from_pdf101(text)
    if auto_sup is not None and bs.get('Szállítók') in (None, 0, ''):
        bs['Szállítók'] = auto_sup
    # --- Final normalization: ensure Szállítók has PREVIOUS ---
    try:
        sup = raw['balance'].get('Szállítók')
        if sup is None or sup.get('previous') in (None,""):
            du_cur, du_prev, du_line = get_trade_payables_dual_universal(text)
            if (du_cur is not None) or (du_prev is not None):
                raw['balance']['Szállítók'] = {'line': du_line, 'current': du_cur, 'previous': du_prev}
                bs['Szállítók'] = du_cur if isinstance(du_cur,(int,float)) else (int(str(du_cur).replace(' ', '')) if du_cur not in (None,'') else None)
    except Exception:
        pass
    # --- /Final normalization ---

    return bs, plv, raw

def use_pick_year_columns_marker(headers):
    # Helper to ensure pick_year_columns is used; keep compatibility
    try:
        return pick_year_columns(headers)
    except Exception:
        return (0, max(0, len(headers)-1))



def compute_ratios(bs, pl):
    """Return core KPIs computed from current-year values. Robust to None/tuples/strings."""
    def to_num(x):
        if isinstance(x, (int, float)):
            return x
        if isinstance(x, str):
            import re as _re
            s = _re.sub(r'[^0-9\-]', '', x)
            try:
                return int(s) if s not in ('', '-') else None
            except Exception:
                return None
        return None

    def safe_div(a, b):
        a = to_num(a); b = to_num(b)
        if a in (None, 0) or b in (None, 0):
            return None
        try:
            return a / b
        except Exception:
            return None

    # Normalize inputs
    CA  = to_num(bs.get('Forgóeszközök'))
    CL  = to_num(bs.get('Rövid lejáratú kötelezettségek'))
    INV = to_num(bs.get('Készletek'))
    TL  = to_num(bs.get('Kötelezettségek összesen'))
    EQ  = to_num(bs.get('Saját tőke'))
    REC = to_num(bs.get('Követelések'))
    NS  = to_num(pl.get('Értékesítés nettó árbevétele'))
    MAT = to_num(pl.get('Anyagjellegű ráfordítások'))
    AP  = to_num(bs.get('Szállítók'))

    # KPIs
    cr = safe_div(CA, CL)

    qr = None
    if CA is not None and INV is not None and CL not in (None, 0):
        qr = safe_div(CA - INV, CL)

    dte = safe_div(TL, EQ)

    nwc = None
    if CA is not None and CL is not None:
        nwc = CA - CL

    rcv_days = None
    if NS not in (None, 0) and REC is not None:
        rcv_days = safe_div(REC, NS)
        if rcv_days is not None:
            rcv_days *= 365

    inv_days = None
    if MAT not in (None, 0) and INV is not None:
        inv_days = safe_div(INV, MAT)
        if inv_days is not None:
            inv_days *= 365

    ap_days = None
    if MAT not in (None, 0) and AP is not None:
        ap_days = safe_div(AP, MAT)
        if ap_days is not None:
            ap_days *= 365

    # Simple scoring (kept close to previous thresholds)
    score = 0
    if cr is not None:
        if cr >= 1.5: score += 20
        elif cr >= 1.2: score += 12
        elif cr >= 1.0: score += 8
        elif cr >= 0.8: score += 5

    if qr is not None:
        if qr >= 1.0: score += 20
        elif qr >= 0.7: score += 12
        elif qr >= 0.5: score += 6

    if dte is not None:
        if dte < 0.5: score += 20
        elif dte < 1.0: score += 14
        elif dte < 2.0: score += 8
        elif dte < 3.0: score += 4

    if rcv_days is not None:
        if rcv_days <= 45: score += 20
        elif rcv_days <= 60: score += 14
        elif rcv_days <= 90: score += 8
        elif rcv_days <= 120: score += 4

    if inv_days is not None:
        if inv_days <= 60: score += 20
        elif inv_days <= 90: score += 14
        elif inv_days <= 120: score += 8
        elif inv_days <= 150: score += 4

    if ap_days is not None:
        if ap_days >= 45: score += 20
        elif ap_days >= 30: score += 14
        elif ap_days >= 20: score += 8
        elif ap_days >= 10: score += 4

    rating = "Excellent" if score >= 85 else "Good" if score >= 70 else "Moderate" if score >= 55 else "Weak"

    return {
        "Current ratio": cr,
        "Quick ratio": qr,
        "Debt/Equity": dte,
        "Nettó forgótőke (eFt)": nwc,
        "Vevőállomány forgási ideje (nap)": rcv_days,
        "Készlet forgási ideje (nap)": inv_days,
        "Szállítói napok (DPO)": ap_days,
        "Kockázati pontszám (0-100)": score,
        "Kockázati besorolás": rating
    }
# ==== Configurable scoring helpers (no UI change) ====
import json, os

def _default_scoring_config():
    return {
        "weights": {"liquidity":25, "wc_cycle":20, "leverage":25, "profitability":30},
        "guards": {
            "loss_one":12, "loss_both":18,
            "ccc_bonus":[[0,-4],[60,0],[120,3],[99999,6]],
            "size_bonus_eFt":[[0,0],[1_000_000,2],[10_000_000,4],[50_000_000,6]]
        },
        "bands": {
            "ebit_margin":[[0,12],[3,8],[6,4],[10,0],[20,-3],[999,-6]],
            "net_margin":[[0,8],[2,4],[5,0],[10,-2],[999,-4]],
            "de":[[0.6,-2],[1.2,0],[1.6,3],[2.0,6],[999,10]],
            "nd_ebitda":[[1.0,-3],[2.0,0],[3.0,3],[5.0,6],[999,10]],
            "ic":[[1.0,12],[2.0,8],[4.0,4],[8.0,0],[999,-3]]
        },
        "sector_benchmarks": {
            "trade":{"dso":45,"dio":60,"dpo":35,"ccc":45},
            "manufacturing":{"dso":60,"dio":90,"dpo":40,"ccc":60},
            "construction":{"dso":75,"dio":75,"dpo":45,"ccc":75},
            "energy":{"dso":75,"dio":75,"dpo":45,"ccc":75}
        }
    }

def load_scoring_config():
    cfg_path = os.path.join(os.path.dirname(__file__), "scoring_config.json")
    try:
        with open(cfg_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return _default_scoring_config()

def _band_points(bands, value):
    if value is None:
        return 0
    v = float(value); last = 0
    for thr, pts in bands:
        if v <= thr:
            return pts
        last = pts
    return last

def sector_key_from_text(sector_text):
    if not sector_text: return None
    s = str(sector_text).lower()
    if "keresk" in s: return "trade"
    if "gyárt" in s or "gyart" in s or "ipar" in s: return "manufacturing"
    if "építő" in s or "epit" in s: return "construction"
    if "energia" in s: return "energy"
    return None

def safe_div(a,b):
    try:
        if a is None or b is None: return None
        b = float(b)
        if b == 0: return None
        return float(a)/b
    except Exception:
        return None

def score_from_rules(ratios, bs, pl, derived, sector_text):
    cfg = load_scoring_config()
    base = 50.0
    # Pull key PL/BS
    revenue = pl.get("Értékesítés nettó árbevétele") or pl.get("Revenue") or 0
    ebit = pl.get("Üzemi (üzleti) tevékenység eredménye")
    netp = pl.get("Adózott eredmény")
    da = pl.get("Értékcsökkenési leírás")
    ebitda = (ebit + da) if (isinstance(ebit,(int,float)) and isinstance(da,(int,float))) else None
    interest = pl.get("Fizetett kamat") or pl.get("Pénzügyi műveletek ráfordításai")
    ic = safe_div(ebit, interest) if (interest and interest>0) else None
    # Net debt
    cash = bs.get("Pénzeszközök"); st = bs.get("Rövid lejáratú kötelezettségek"); lt = bs.get("Hosszú lejáratú kötelezettségek")
    nd = None
    if isinstance(st,(int,float)) or isinstance(lt,(int,float)) or isinstance(cash,(int,float)):
        nd = (st or 0) + (lt or 0) - (cash or 0)
    nd_eb = safe_div(nd, ebitda) if (nd is not None and ebitda and ebitda>0) else None
    ebit_m = safe_div(ebit, revenue); net_m = safe_div(netp, revenue)

    # Points from bands
    pts = 0.0; bands = cfg["bands"]
    if ebit_m is not None: pts += _band_points(bands["ebit_margin"], ebit_m*100.0)
    if net_m  is not None: pts += _band_points(bands["net_margin"],  net_m*100.0)
    if derived.get("de")  is not None: pts += _band_points(bands["de"], float(derived["de"]))
    if nd_eb is not None: pts += _band_points(bands["nd_ebitda"], nd_eb)
    if ic    is not None: pts += _band_points(bands["ic"], ic)

    # Liquidity nudges (compatible with UI)
    cr = derived.get("cr"); qr = derived.get("qr")
    if cr is not None: pts += (4 if cr < (ratios.get("CR_MIN") or 1.2) else (-2 if cr > (ratios.get("CR_GOOD") or 1.5) else 0))
    if qr is not None: pts += (4 if qr < (ratios.get("QR_MIN") or 1.0) else (-2 if qr > (ratios.get("QR_GOOD") or 1.2) else 0))

    # Working capital cycle vs sector
    sec_key = sector_key_from_text(sector_text)
    sb = cfg["sector_benchmarks"].get(sec_key, {"dso":60,"dio":90,"dpo":40,"ccc":60})
    dso = derived.get("dso"); dio = derived.get("dio"); dpo = derived.get("dpo"); ccc = derived.get("ccc")
    if dso is not None: pts += (3 if dso > sb["dso"] else 0)
    if dio is not None: pts += (3 if dio > sb["dio"] else 0)
    if dpo is not None: pts += (3 if dpo < sb["dpo"] else 0)
    if ccc is not None:
        for thr, add in cfg["guards"]["ccc_bonus"]:
            if ccc <= thr: pts += add; break

    # Guards
    if isinstance(ebit,(int,float)) and ebit < 0 and isinstance(netp,(int,float)) and netp < 0:
        pts += cfg["guards"]["loss_both"]
    elif (isinstance(ebit,(int,float)) and ebit < 0) or (isinstance(netp,(int,float)) and netp < 0):
        pts += cfg["guards"]["loss_one"]

    # Size bonus (revenue in eFt)
    rev_eFt = revenue if isinstance(revenue,(int,float)) else 0
    last_bonus = 0
    for thr, bonus in cfg["guards"]["size_bonus_eFt"]:
        if rev_eFt <= thr: pts -= bonus; break
        last_bonus = bonus
    else:
        pts -= last_bonus

    score = base + pts
    if score < 0: score = 0.0
    if score > 100: score = 100.0
    return round(score,1)
# ==== End scoring helpers ====
def _safe_to_num(x):
    try:
        if x is None: return None
        if isinstance(x,(int,float)): return float(x)
        # strip spaces and non-breaking spaces, thousands separators
        s = str(x).replace("\xa0"," ").replace(" "," ").replace(" "," ").strip()
        s = s.replace(".", " ").replace(",", " ").replace("\u2212","-").replace("\u2012","-").replace("\u2013","-").replace("\u2014","-")
        s = " ".join(s.split())
        toks = [t for t in s.split(" ") if any(ch.isdigit() for ch in t) or t in ("-","+")]
        if not toks: return None
        s2 = "".join(toks)
        return float(int(s2))
    except Exception:
        return None

def _first_present(dct, keys):
    if not isinstance(dct, dict): return None
    for k in keys:
        if k in dct and _safe_to_num(dct.get(k)) is not None:
            return _safe_to_num(dct.get(k))
    return None

def build_cf_section(doc, lang_code: str, bs_curr, bs_prev, pl_curr):
    from docx.enum.table import WD_TABLE_ALIGNMENT
    try:
        def _num(x):
            try:
                if x is None: return None
                if isinstance(x,(int,float)): return float(x)
                s = str(x).replace("\xa0","").replace(" ","").replace(",","").replace(".","")
                if s in ("", "-", "—"): return None
                return float(int(s))
            except: return None
        def _first(dct, keys):
            if not isinstance(dct, dict): return None
            for k in keys:
                if k in dct:
                    vn = _num(dct.get(k))
                    if vn is not None: return vn
            return None
        # Picks
        NI = _first(pl_curr, ["Adózott eredmény","Profit after tax","Net income","Net profit"])
        DA = _first(pl_curr, ["Értékcsökkenési leírás","ÉCS","Depreciation and amortization","Depreciation","Amortization"])
        ARc = _first(bs_curr, ["Vevők","Vevőkövetelések","Követelések","Receivables","Trade receivables"])
        ARp = _first(bs_prev, ["Vevők","Vevőkövetelések","Követelések","Receivables","Trade receivables"])
        INVc = _first(bs_curr, ["Készletek","Készlet","Inventory","Inventories"])
        INVp = _first(bs_prev, ["Készletek","Készlet","Inventory","Inventories"])
        APc  = _first(bs_curr, ["Szállítók","Kötelezettségek - Szállítók","Payables","Trade payables"])
        APp  = _first(bs_prev, ["Szállítók","Kötelezettségek - Szállítók","Payables","Trade payables"])
        CASHc = _first(bs_curr, ["Pénzeszközök","Cash and cash equivalents"])
        CASHp = _first(bs_prev, ["Pénzeszközök","Cash and cash equivalents"])
        STLc  = _first(bs_curr, ["Rövid lejáratú kötelezettségek","Short-term liabilities"])
        STLp  = _first(bs_prev, ["Rövid lejáratú kötelezettségek","Short-term liabilities"])
        LTLc  = _first(bs_curr, ["Hosszú lejáratú kötelezettségek","Long-term liabilities"])
        LTLp  = _first(bs_prev, ["Hosszú lejáratú kötelezettségek","Long-term liabilities"])
        EQc   = _first(bs_curr, ["Saját tőke","Equity"])
        EQp   = _first(bs_prev, ["Saját tőke","Equity"])
        OCAc = _first(bs_curr, ["Egyéb követelések","Egyéb rövid lejáratú követelések","Aktív időbeli elhatárolások","Other receivables","Other current assets","Prepayments and accrued income"])
        OCAp = _first(bs_prev, ["Egyéb követelések","Egyéb rövid lejáratú követelések","Aktív időbeli elhatárolások","Other receivables","Other current assets","Prepayments and accrued income"])
        OCLc = _first(bs_curr, ["Egyéb rövid lejáratú kötelezettségek","Passzív időbeli elhatárolások","Other current liabilities","Accrued expenses and deferred income"])
        OCLp = _first(bs_prev, ["Egyéb rövid lejáratú kötelezettségek","Passzív időbeli elhatárolások","Other current liabilities","Accrued expenses and deferred income"])
        FA_c = _first(bs_curr, ["Befektetett eszközök","Tárgyi eszközök","Immateriális javak","Fixed assets","Property, plant and equipment","Intangible assets"])
        FA_p = _first(bs_prev, ["Befektetett eszközök","Tárgyi eszközök","Immateriális javak","Fixed assets","Property, plant and equipment","Intangible assets"])
        # Deltas
        dAR   = None if (ARc is None or ARp is None) else (ARc - ARp)
        dINV  = None if (INVc is None or INVp is None) else (INVc - INVp)
        dAP   = None if (APc is None or APp is None) else (APc - APp)
        dOCA  = None if (OCAc is None or OCAp is None) else (OCAc - OCAp)
        dOCL  = None if (OCLc is None or OCLp is None) else (OCLc - OCLp)
        dNWC  = None if (dAR is None or dINV is None or dAP is None) else (dAR + dINV - dAP)
        dCASH = None if (CASHc is None or CASHp is None) else (CASHc - CASHp)
        dSTL  = None if (STLc is None or STLp is None) else (STLc - STLp)
        dLTL  = None if (LTLc is None or LTLp is None) else (LTLc - LTLp)
        dEQ   = None if (EQc is None or EQp is None) else (EQc - EQp)
        dSTL_exAP = None if (dSTL is None or dAP is None) else (dSTL - dAP)
        dFA   = None if (FA_c is None or FA_p is None) else (FA_c - FA_p)
        CFO = None if (NI is None and DA is None and dNWC is None) else ((NI or 0) + (DA or 0) - (0 if dNWC is None else dNWC))
        CFI_CFF_proxy = None if (dCASH is None or CFO is None) else (dCASH - CFO)
        # Render
        def add_table(lang):
            tb = doc.add_table(rows=1, cols=2); tb.alignment = WD_TABLE_ALIGNMENT.LEFT; tb.style = "Light Grid Accent 1"
            if lang=="hu": tb.cell(0,0).text, tb.cell(0,1).text = "Tétel","Összeg (eFt)"
            else:          tb.cell(0,0).text, tb.cell(0,1).text = "Item","Amount (th HUF)"
            def put(label, val):
                r = tb.add_row().cells
                r[0].text = label
                r[1].text = ("n.a." if val is None else (f"{int(round(val)):,}".replace(","," ") if lang=="hu" else f"{int(round(val)):,}"))
            return put
        if lang_code=="hu":
            doc.add_heading("5) Pénzáram (Cash-flow) – részletes", level=1)
            put = add_table("hu")
            put("Vevőkövetelések változása (−ΔAR)", (None if dAR is None else -dAR))
            put("Készletek változása (−ΔKészlet)", (None if dINV is None else -dINV))
            put("Szállítói kötelezettségek változása (+ΔSzállítók)", dAP)
            if dOCA is not None: put("Egyéb rövid lejáratú eszközök változása (−ΔOCA)", -dOCA)
            if dOCL is not None: put("Egyéb rövid lejáratú kötelezettségek változása (+ΔOCL)", dOCL)
            put("Nettó forgótőke változása (ΔNWC)", dNWC)
            put("Működési cash-flow (CFO, becslés)", CFO)
            put("Pénzeszközök változása (ΔCash)", dCASH)
            put("Rövid lejáratú kötelezettségek változása (ΔSTL)", dSTL)
            put("ebből: ΔSTL − ΔSzállítók (adósság proxy)", dSTL_exAP)
            put("Hosszú lejáratú kötelezettségek változása (ΔLTL)", dLTL)
            put("Saját tőke változása (ΔEquity)", dEQ)
            if dFA is not None: put("Befektetett eszközök változása (ΔFixed assets, proxy)", dFA)
            put("Nettó nem-működési CF (CFI+CFF proxy) = ΔCash − CFO", CFI_CFF_proxy)
            parts = []
            def sgn(x): 
                return None if x is None else (1 if x>0 else (-1 if x<0 else 0))
            if sgn(dAR)==1:   parts.append(f"A vevők {int(abs(dAR)):,} eFt-tal nőttek → készpénz LEKÖTÉS.".replace(","," "))
            if sgn(dAR)==-1:  parts.append(f"A vevők {int(abs(dAR)):,} eFt-tal csökkentek → készpénz BEÁRAMLÁS.".replace(","," "))
            if sgn(dINV)==1:  parts.append(f"A készlet {int(abs(dINV)):,} eFt-tal nőtt → készpénz LEKÖTÉS.".replace(","," "))
            if sgn(dINV)==-1: parts.append(f"A készlet {int(abs(dINV)):,} eFt-tal csökkent → készpénz BEÁRAMLÁS.".replace(","," "))
            if sgn(dAP)==1:   parts.append(f"A szállítók {int(abs(dAP)):,} eFt-tal nőttek → készpénz MEGTARTÁS (későbbi fizetés).".replace(","," "))
            if sgn(dAP)==-1:  parts.append(f"A szállítók {int(abs(dAP)):,} eFt-tal csökkentek → készpénz KIÁRAMLÁS (gyorsabb fizetés).".replace(","," "))
            if sgn(dOCA)==1:  parts.append(f"Egyéb rövid lej. eszközök +{int(abs(dOCA)):,} eFt → készpénz LEKÖTÉS.".replace(","," "))
            if sgn(dOCA)==-1: parts.append(f"Egyéb rövid lej. eszközök −{int(abs(dOCA)):,} eFt → készpénz BEÁRAMLÁS.".replace(","," "))
            if sgn(dOCL)==1:  parts.append(f"Egyéb rövid lej. kötelezettségek +{int(abs(dOCL)):,} eFt → készpénz MEGTARTÁS.".replace(","," "))
            if sgn(dOCL)==-1: parts.append(f"Egyéb rövid lej. kötelezettségek −{int(abs(dOCL)):,} eFt → készpénz KIÁRAMLÁS.".replace(","," "))
            if dNWC is not None:
                parts.append(("ΔNWC +{:,} eFt → nettó készpénz LEKÖTÉS a működésben." if dNWC>0 else "ΔNWC {:,} eFt → nettó készpénz FELSZABADULÁS a működésből.").format(int(dNWC)).replace(","," "))
            if CFO is not None:
                parts.append(("Eredő CFO (becslés): +{:,} eFt → a működés pénzt termel." if CFO>=0 else "Eredő CFO (becslés): {:,} eFt → a működés pénzt éget.").format(int(CFO)).replace(","," "))
            if CFI_CFF_proxy is not None:
                if CFI_CFF_proxy < 0:
                    msg = f"Nem-működési CF (CFI+CFF proxy): {int(CFI_CFF_proxy):,} eFt → valószínűleg BERUHÁZÁS/ADÓSSÁGSZOLGÁLAT/OSZTALÉK kifizetés."
                else:
                    msg = f"Nem-működési CF (CFI+CFF proxy): +{int(CFI_CFF_proxy):,} eFt → külső forrás BEÁRAMLÁS (hitelfelvétel/tőkeinjekció/értékesítés)."
                parts.append(msg.replace(","," "))
            if sgn(dLTL)==1 or sgn(dSTL_exAP)==1: parts.append("Finanszírozási jel: adósságállomány nőtt (ΔLTL/ΔSTL_exAP > 0)")
            if sgn(dEQ)==1: parts.append("Tőkejel: saját tőke nőtt → lehetséges tőkeinjekció")
            if sgn(dFA)==1 and CFI_CFF_proxy is not None and CFI_CFF_proxy<0:
                parts.append("Befektetés-jel: a befektetett eszközök állománya nőtt, és a nem-működési CF negatív → nagy valószínűséggel beruházás történt.")
            for t in parts: doc.add_paragraph("• " + t)

            # --- CF KPI mini-block (HU) ---
            try:
                rev = _first(pl_curr, ["Értékesítés nettó árbevétele","Net sales revenue","Sales"])
                def pct(x, base):
                    if x is None or base in (None,0): return None
                    return 100.0 * float(x) / float(base)
                CFO_margin = None if (CFO is None) else pct(CFO, rev)
                FCF_proxy = None if CFO is None else (CFO - max(dFA or 0, 0))
                FCF_margin = None if (FCF_proxy is None) else pct(FCF_proxy, rev)
                NWC_int = None if dNWC is None else pct(dNWC, rev)
                mat = _first(pl_curr, ["Anyagjellegű ráfordítások","Material-type expenses"])
                pers = _first(pl_curr, ["Személyi jellegű ráfordítások","Personnel expenses"])
                oth = _first(pl_curr, ["Egyéb ráfordítások","Other expenses"])
                ecs = DA
                monthly_burn = None
                if any(v is not None for v in [mat,pers,oth]):
                    base = (mat or 0)+(pers or 0)+(oth or 0)-(ecs or 0)
                    monthly_burn = base/12.0 if base and base>0 else None
                runway = None if (CASHc is None or monthly_burn in (None,0)) else (CASHc/monthly_burn)
                def band(v, bands):
                    if v is None: return 50
                    for thr, pts in bands:
                        if v <= thr: return pts
                    return bands[-1][1]
                CFO_pts = band(-(CFO_margin or -9999), [[-9999,10],[-10,30],[-5,55],[0,70],[5,85],[10,95],[9999,100]])
                FCF_pts = band(-(FCF_margin or -9999), [[-9999,10],[-10,30],[-5,55],[0,70],[5,85],[9999,95]])
                NWC_pts = band(abs(NWC_int or 0), [[0,100],[3,85],[8,70],[999,50]])
                RW_pts = band(runway or 0, [[1,30],[3,55],[6,80],[999,95]])
                debt_flag = 1 if ((dSTL_exAP and dSTL_exAP>0) or (dLTL and dLTL>0)) and (CFO is not None and CFO<=0) else 0
                penalty = 10 if debt_flag else 0
                CF_score = max(0, min(100, round((0.25*CFO_pts + 0.25*FCF_pts + 0.15*NWC_pts + 0.15*RW_pts + 0.10*(100 if (dFA or 0)<=0 else 70) + 0.10*(100-penalty)), 0)))
                tb2 = doc.add_table(rows=1, cols=3); tb2.alignment=WD_TABLE_ALIGNMENT.LEFT; tb2.style="Light Grid Accent 1"
                tb2.cell(0,0).text="Mutató"; tb2.cell(0,1).text="Érték"; tb2.cell(0,2).text="Megjegyzés"
                def row(name, val, note=""):
                    r=tb2.add_row().cells
                    r[0].text=name
                    r[1].text = ("n.a." if val is None else ( (f"{str(val).replace('.',',')}" if "futamidő" in name.lower() else f"{str(round(val,1)).replace('.',',')}%") if isinstance(val,float) else f"{int(round(val)):,}".replace(","," ") ) )
                    r[2].text = note
                row("CFO margin", CFO_margin, "")
                row("FCF (proxy) margin", FCF_margin, "")
                row("ΔNWC / Árbevétel", NWC_int, "")
                row("Likviditási futamidő (hó)", runway, "")
                doc.add_paragraph(f"CF minősítés: {int(CF_score)}/100")
                globals()['CF_SCORE_LAST'] = CF_score
            except Exception:
                pass
        else:
            doc.add_heading("5) Cash-flow – detailed", level=1)
            put = add_table("en")
            put("Change in receivables (−ΔAR)", (None if dAR is None else -dAR))
            put("Change in inventory (−ΔINV)", (None if dINV is None else -dINV))
            put("Change in trade payables (+ΔAP)", dAP)
            if dOCA is not None: put("Change in other current assets (−ΔOCA)", -dOCA)
            if dOCL is not None: put("Change in other current liabilities (+ΔOCL)", dOCL)
            put("Net working capital change (ΔNWC)", dNWC)
            put("Operating cash flow (CFO, est.)", CFO)
            put("Change in cash (ΔCash)", dCASH)
            put("Change in short-term liabilities (ΔSTL)", dSTL)
            put("of which: ΔSTL − ΔAP (debt proxy)", dSTL_exAP)
            put("Change in long-term liabilities (ΔLTL)", dLTL)
            put("Change in equity (ΔEquity)", dEQ)
            if dFA is not None: put("Change in fixed assets (ΔFA, proxy)", dFA)
            put("Net non-operational CF (CFI+CFF proxy) = ΔCash − CFO", CFI_CFF_proxy)
            parts = []
            def sgn(x): 
                return None if x is None else (1 if x>0 else (-1 if x<0 else 0))
            if sgn(dAR)==1:   parts.append(f"Receivables up by {int(abs(dAR)):,} → CASH TIED UP.")
            if sgn(dAR)==-1:  parts.append(f"Receivables down by {int(abs(dAR)):,} → CASH INFLOW.")
            if sgn(dINV)==1:  parts.append(f"Inventory up by {int(abs(dINV)):,} → CASH TIED UP.")
            if sgn(dINV)==-1: parts.append(f"Inventory down by {int(abs(dINV)):,} → CASH INFLOW.")
            if sgn(dAP)==1:   parts.append(f"Payables up by {int(abs(dAP)):,} → CASH PRESERVED (later payment).")
            if sgn(dAP)==-1:  parts.append(f"Payables down by {int(abs(dAP)):,} → CASH OUTFLOW (faster payment).")
            if sgn(dOCA)==1:  parts.append(f"Other current assets +{int(abs(dOCA)):,} → CASH TIED UP.")
            if sgn(dOCA)==-1: parts.append(f"Other current assets −{int(abs(dOCA)):,} → CASH INFLOW.")
            if sgn(dOCL)==1:  parts.append(f"Other current liabilities +{int(abs(dOCL)):,} → CASH PRESERVED.")
            if sgn(dOCL)==-1: parts.append(f"Other current liabilities −{int(abs(dOCL)):,} → CASH OUTFLOW.")
            if dNWC is not None:
                parts.append(("ΔNWC +{:,} → net CASH TIED UP in operations." if dNWC>0 else "ΔNWC {:,} → net CASH RELEASED from operations.").format(int(dNWC)))
            if CFO is not None:
                parts.append(("Resulting CFO (est.): +{:,} → operations GENERATE cash." if CFO>=0 else "Resulting CFO (est.): {:,} → operations CONSUME cash.").format(int(CFO)))
            if CFI_CFF_proxy is not None:
                parts.append(("Non-operational CF (CFI+CFF proxy): {:,} → likely INVESTMENT/DEBT SERVICE/DIVIDEND outflow." if CFI_CFF_proxy<0 else "Non-operational CF (CFI+CFF proxy): +{:,} → external inflow (new debt/equity/asset disposal).").format(int(CFI_CFF_proxy)))
            if sgn(dLTL)==1 or sgn(dSTL_exAP)==1: parts.append("Financing signal: debt load increased (ΔLTL/ΔSTL_exAP > 0)")
            if sgn(dEQ)==1: parts.append("Equity signal: equity increased → possible injection")
            if sgn(dFA)==1 and CFI_CFF_proxy is not None and CFI_CFF_proxy<0:
                parts.append("Investment signal: fixed assets increased and non-operational CF is negative → likely capex.")
            for t in parts: doc.add_paragraph("• " + t)

            # --- CF KPI mini-block (EN) ---
            try:
                rev = _first(pl_curr, ["Értékesítés nettó árbevétele","Net sales revenue","Sales"])
                def pct(x, base):
                    if x is None or base in (None,0): return None
                    return 100.0 * float(x) / float(base)
                CFO_margin = None if (CFO is None) else pct(CFO, rev)
                FCF_proxy = None if CFO is None else (CFO - max(dFA or 0, 0))
                FCF_margin = None if (FCF_proxy is None) else pct(FCF_proxy, rev)
                NWC_int = None if dNWC is None else pct(dNWC, rev)
                mat = _first(pl_curr, ["Anyagjellegű ráfordítások","Material-type expenses"])
                pers = _first(pl_curr, ["Személyi jellegű ráfordítások","Personnel expenses"])
                oth = _first(pl_curr, ["Egyéb ráfordítások","Other expenses"])
                ecs = DA
                monthly_burn = None
                if any(v is not None for v in [mat,pers,oth]):
                    base = (mat or 0)+(pers or 0)+(oth or 0)-(ecs or 0)
                    monthly_burn = base/12.0 if base and base>0 else None
                runway = None if (CASHc is None or monthly_burn in (None,0)) else (CASHc/monthly_burn)
                def band(v, bands):
                    if v is None: return 50
                    for thr, pts in bands:
                        if v <= thr: return pts
                    return bands[-1][1]
                CFO_pts = band(-(CFO_margin or -9999), [[-9999,10],[-10,30],[-5,55],[0,70],[5,85],[10,95],[9999,100]])
                FCF_pts = band(-(FCF_margin or -9999), [[-9999,10],[-10,30],[-5,55],[0,70],[5,85],[9999,95]])
                NWC_pts = band(abs(NWC_int or 0), [[0,100],[3,85],[8,70],[999,50]])
                RW_pts = band(runway or 0, [[1,30],[3,55],[6,80],[999,95]])
                debt_flag = 1 if ((dSTL_exAP and dSTL_exAP>0) or (dLTL and dLTL>0)) and (CFO is not None and CFO<=0) else 0
                penalty = 10 if debt_flag else 0
                CF_score = max(0, min(100, round((0.25*CFO_pts + 0.25*FCF_pts + 0.15*NWC_pts + 0.15*RW_pts + 0.10*(100 if (dFA or 0)<=0 else 70) + 0.10*(100-penalty)), 0)))
                tb2 = doc.add_table(rows=1, cols=3); tb2.alignment=WD_TABLE_ALIGNMENT.LEFT; tb2.style="Light Grid Accent 1"
                tb2.cell(0,0).text="Metric"; tb2.cell(0,1).text="Value"; tb2.cell(0,2).text="Note"
                def row(name, val, note=""):
                    r=tb2.add_row().cells
                    r[0].text=name
                    r[1].text = ("n.a." if val is None else ( (f"{val:.1f}" if "runway" in name.lower() else f"{val:.1f}%") if isinstance(val,float) else f"{int(round(val)):,}") )
                    r[2].text = note
                row("CFO margin", CFO_margin, "")
                row("FCF (proxy) margin", FCF_margin, "")
                row("ΔNWC / Revenue", NWC_int, "")
                row("Liquidity runway (months)", runway, "")
                doc.add_paragraph(f"CF rating: {int(CF_score)}/100")
                globals()['CF_SCORE_LAST'] = CF_score
            except Exception:
                pass
        return True
    except Exception:
        try:
            if lang_code=="hu":
                doc.add_heading("5) Pénzáram (Cash-flow)", level=1); doc.add_paragraph("CF suppressed – hiba a CF modulban.")
            else:
                doc.add_heading("5) Cash-flow", level=1); doc.add_paragraph("CF suppressed – error in CF module.")
        except: pass
        return False

def make_docx(company_name, bs, pl, ratios, out_path: Path, sector='default', lang='hu', prev=None, raw=None):
    # Final design as agreed (HU/EN mirror, merged 2/a+3/b table, per-KPI method rows)
    from pathlib import Path as _Path
    from docx import Document
    from docx.shared import RGBColor, Pt
    from docx.oxml.ns import qn
    from docx.enum.table import WD_TABLE_ALIGNMENT


    def _reparse_from_raw_line(line):
        try:
            # fallback simple: numbers in line -> first is prev, last is current
            import re
            nums = _extract_grouped_numbers(line or '')
            def _to_int(tok):
                if not tok: return None
                neg = tok.strip().startswith('(') and tok.strip().endswith(')')
                tok = tok.strip().strip('()').replace(' ', '')
                try:
                    v = int(tok)
                    return -v if neg else v
                except:
                    return None
            if not nums:
                return None, None
            pv = nums[0] if isinstance(nums[0], int) else None; cv = nums[-1] if isinstance(nums[-1], int) else None
            return pv, cv
        except:
            return None, None

    # ---- config / benchmarks ----
    cfg_path = Path(__file__).with_name('benchmarks.json')
    try:
        _cfg_all = json.loads(cfg_path.read_text(encoding='utf-8'))
    except Exception:
        _cfg_all = {
            "default": {
                "targets": {
                    "current_ratio": 1.2, "quick_ratio": 1.0, "debt_to_equity": 1.6,
                    "receivables_days": 60, "inventory_days": 90, "payables_days_min": 40
                }
            }
        }
    T = _cfg_all.get(sector, _cfg_all.get("default", {})).get("targets", {})
    DF = _cfg_all.get("default", {}).get("targets", {})
    CR_MIN = float(T.get('current_ratio', DF.get('current_ratio', 1.2)))
    QR_MIN = float(T.get('quick_ratio', DF.get('quick_ratio', 1.0)))
    DE_MAX = float(T.get('debt_to_equity', DF.get('debt_to_equity', 1.6)))
    DSO_MAX = float(T.get('receivables_days', DF.get('receivables_days', 60)))
    DIO_MAX = float(T.get('inventory_days', DF.get('inventory_days', 90)))
    DPO_MIN = float(T.get('payables_days_min', DF.get('payables_days_min', 40)))

    # ---- helpers ----
    def badge(run, status):
        run.text = "●"
        col = {"green": RGBColor(22,163,74), "yellow": RGBColor(202,138,4), "red": RGBColor(220,38,38)}.get(status, RGBColor(107,114,128))
        run.font.color.rgb = col; run.bold = True

    def status_of(value, kind):
        if value is None or not isinstance(value,(int,float)): return "yellow"
        if kind=="current": return "green" if value>=CR_MIN else "red"
        if kind=="quick":   return "green" if value>=QR_MIN else "red"
        if kind=="de":      return "green" if value<=DE_MAX else "red"
        if kind=="dso":     return "green" if value<=DSO_MAX else "red"
        if kind=="dio":     return "green" if value<=DIO_MAX else "red"
        if kind=="dpo":     return "green" if value>=DPO_MIN else "red"
        if kind=="ccc":
            v=value
            return "red" if v>120 else ("yellow" if v>60 else "green")
        return "yellow"

    def fmt_num_hu(v, kind):
        try: v=float(v)
        except: return "n.a."
        if kind=="times": return f"{v:.2f}×".replace(".", ",")
        if kind=="days":  return f"{v:.1f} nap".replace(".", ",")
        return f"{v:.2f}".replace(".", ",")

    def fmt_num_en(v, kind):
        try: v=float(v)
        except: return "n.a."
        if kind=="times": return f"{v:.2f}×"
        if kind=="days":  return f"{v:.1f} d"
        return f"{v:.2f}"

    def range_text(color, lang_code):
        if lang_code=="hu":
            return "Alacsony = 0–39 pont" if color=="green" else ("Közepes = 40–69 pont" if color=="yellow" else "Magas = 70–100 pont")
        else:
            return "Low = 0–39" if color=="green" else ("Moderate = 40–69" if color=="yellow" else "High = 70–100")

    # extract metrics
    cr  = ratios.get("Current ratio")
    qr  = ratios.get("Quick ratio")
    dte = ratios.get("Debt/Equity")
    dso = ratios.get("Vevőállomány forgási ideje (nap)")
    dio = ratios.get("Készlet forgási ideje (nap)")
    dpo = ratios.get("Szállítói napok (DPO)")
    ccc = (dso + dio - dpo) if all(isinstance(x,(int,float)) for x in (dso,dio,dpo)) else ratios.get("CCC")
    try:
        _recv = bs.get('Követelések') or 0
        _inv  = bs.get('Készletek') or 0
        _pay  = bs.get('Szállítók') or 0
        _curr = bs.get('Forgóeszközök') or 0
        _stl  = bs.get('Rövid lejáratú kötelezettségek') or 0
        wcn = (_recv + _inv - _pay) if any(isinstance(x,(int,float)) for x in (_recv,_inv,_pay)) else None
        nwc = (_curr - _stl) if any(isinstance(x,(int,float)) for x in (_curr,_stl)) else None
    except Exception:
        wcn = None; nwc = None

    derived = {"cr":cr,"qr":qr,"de":dte,"dso":dso,"dio":dio,"dpo":dpo,"ccc":ccc}
    try:
        score = score_from_rules(ratios, bs, pl, derived, sector)
    except Exception:
        score = ratios.get("Kockázati pontszám (0-100)")

    # Integrate CF score at 20% weight if available
    try:
        cf_local = globals().get('CF_SCORE_LAST', None)
        if isinstance(score,(int,float)) and isinstance(cf_local,(int,float)):
            score = round(0.8*float(score) + 0.2*float(cf_local), 1)
    except Exception:
        pass
    rating = (ratios.get("Kockázati besorolás") or "").lower()
    color = "green" if isinstance(score,(int,float)) and score <= 39 else ("yellow" if isinstance(score,(int,float)) and score <= 69 else ("red" if isinstance(score,(int,float)) else ("green" if ("alacsony" in rating or "low" in rating) else ("yellow" if ("közepes" in rating or "moderate" in rating) else ("red" if rating else "yellow")))))

    # interpretations
    def interp(kind, lang_code):
        vmap = {"current":cr,"quick":qr,"de":dte,"dso":dso,"dio":dio,"dpo":dpo,"ccc":ccc}
        v = vmap.get(kind)
        def dec_hu(x):
            try:
                return str(round(float(x),1)).replace('.',',')
            except:
                return "n.a."
        if lang_code=="hu":
            if kind=="current":
                if v is None: return "Nincs adat."
                return f"Minden 1 Ft rövid tartozásra ~{dec_hu(v)} Ft forgóeszköz jut → a napi kiadások fedezettek, váratlan tétel is kezelhető." if v>=CR_MIN else "Fedezet szűk: a rövid tartozásokhoz kevés forgóeszköz társul."
            if kind=="quick":
                if v is None: return "Nincs adat."
                return f"Készletek nélkül is rendben: {str(round(v,2)).replace('.',',')}× fedezet a rövid tartozásokra." if v>=QR_MIN else "Készletek nélkül kevés az azonnali fedezet."
            if kind=="de":
                if v is None: return "Nincs adat."
                return f"Hitelek aránya magas ({str(round(v,2)).replace('.',',')}× vs {str(round(DE_MAX,2)).replace('.',',')}× limit) → érzékenyebb a kamatokra és feltételekre." if v>DE_MAX else "Tőkeáttétel kezelhető tartományban."
            if kind=="dso":
                if v is None: return "Nincs adat."
                diff = v-DSO_MAX
                return f"Vevők átlagosan ~{dec_hu(diff)} nappal később fizetnek → több pénz ragad be követelésekben." if v>DSO_MAX else "Vevői fizetési idő a célszinten belül."
            if kind=="dio":
                if v is None: return "Nincs adat."
                return "Készlet forgása rendben → nem a készlet köti le a pénzt." if v<=DIO_MAX else "Készletnapok magasabbak a szokásosnál → több pénz van készletben."
            if kind=="dpo":
                if v is None: return "Nincs adat."
                diff = DPO_MIN - v
                return f"Túl korán fizetünk (~{dec_hu(diff)} nap) → feleslegesen viszi a készpénzt; érdemes határidőt hosszabbítani." if v<DPO_MIN else "Szállítói napok összhangban az ajánlottal."
            if kind=="ccc":
                if v is None: return "Nincs adat."
                try:
                    months = round(v/30,1)
                    months = str(months).replace('.',',')
                except:
                    months = "≈"
                return f"Hosszú pénzciklus: a működés ~{months} hónapra leköti a pénzt → külső forrás enyhítheti a terhet." if v>120 else ("Közepes ciklus: érdemes figyelni." if v>60 else "Rövid ciklus: hatékony pénzforgás.")
        else:
            if kind=="current":
                if v is None: return "n.a."
                return f"For every 1 of short-term debt there is ~{round(float(v),1)} of current assets → day-to-day bills are well covered." if v>=CR_MIN else "Coverage is thin vs short-term debts."
            if kind=="quick":
                if v is None: return "n.a."
                return f"Even without inventory there is {round(float(v),2)}× cover of short-term debts." if v>=QR_MIN else "Limited instant cover without inventory."
            if kind=="de":
                if v is None: return "n.a."
                return f"Leverage is high ({round(float(v),2)}× vs {round(float(DE_MAX),2)}× cap) → more sensitive to interest and lender terms." if v>DE_MAX else "Leverage within acceptable range."
            if kind=="dso":
                if v is None: return "n.a."
                diff = v-DSO_MAX
                return f"Customers pay ~+{round(diff,1)} d slower → more cash stuck in receivables." if v>DSO_MAX else "Collections within target."
            if kind=="dio":
                if v is None: return "n.a."
                return "Inventory turnover is fine → stock isn’t the cash bottleneck." if v<=DIO_MAX else "Inventory days are elevated → more cash tied up."
            if kind=="dpo":
                if v is None: return "n.a."
                diff = DPO_MIN - v
                return f"We pay suppliers too early (~{round(diff,1)} d) → unnecessary cash drain; extend terms." if v<DPO_MIN else "Payables days in line with guidance."
            if kind=="ccc":
                if v is None: return "n.a."
                months = round(v/30,1)
                return f"Long cash loop: cash is tied up for ~{months} months → extra funding can ease pressure." if v>120 else ("Mid-length cycle: monitor." if v>60 else "Short cycle: efficient.")
        return ""

    # build KPI list
    kpis = [
        {"key":"current","hu":"Current ratio (likviditási ráta)","en":"Current ratio (liquidity)","val":cr,"bmk_hu":f"≥ {CR_MIN:.2f}×".replace('.',','),"bmk_en":f"≥ {CR_MIN:.2f}×",
         "status":status_of(cr,"current"), "method_hu":"Forgóeszközök / Rövid lejáratú kötelezettségek", "method_en":"Current Assets / Short-term Liabilities"},
        {"key":"quick","hu":"Quick ratio (gyorsráta)","en":"Quick ratio (acid-test)","val":qr,"bmk_hu":f"≥ {QR_MIN:.2f}×".replace('.',','),"bmk_en":f"≥ {QR_MIN:.2f}×",
         "status":status_of(qr,"quick"), "method_hu":"(Forgóeszközök − Készletek) / Rövid lejáratú kötelezettségek", "method_en":"(Current Assets − Inventory) / Short-term Liabilities"},
        {"key":"de","hu":"Debt/Equity (tőkeáttétel)","en":"Debt/Equity (leverage)","val":dte,"bmk_hu":f"≤ {DE_MAX:.2f}×".replace('.',','),"bmk_en":f"≤ {DE_MAX:.2f}×",
         "status":status_of(dte,"de"), "method_hu":"Kötelezettségek összesen / Saját tőke", "method_en":"Total Liabilities / Equity"},
        {"key":"dso","hu":"DSO – Vevőnapok","en":"DSO – Days Sales Outstanding","val":dso,"bmk_hu":f"≤ {int(DSO_MAX)} nap","bmk_en":f"≤ {int(DSO_MAX)} d",
         "status":status_of(dso,"dso"), "method_hu":"Vevőkövetelések / Árbevétel × 365", "method_en":"Receivables / Revenue × 365"},
        {"key":"dio","hu":"DIO – Készletnapok","en":"DIO – Days Inventory Outstanding","val":dio,"bmk_hu":f"≤ {int(DIO_MAX)} nap","bmk_en":f"≤ {int(DIO_MAX)} d",
         "status":status_of(dio,"dio"), "method_hu":"Készletek / Anyagjellegű ráfordítások × 365", "method_en":"Inventory / COGS × 365"},
        {"key":"dpo","hu":"DPO – Szállítói napok","en":"DPO – Days Payables Outstanding","val":dpo,"bmk_hu":f"≥ {int(DPO_MIN)} nap","bmk_en":f"≥ {int(DPO_MIN)} d",
         "status":status_of(dpo,"dpo"), "method_hu":"Szállítók / Anyagjellegű ráfordítások × 365", "method_en":"Payables / COGS × 365"},
    ]
    if isinstance(ccc,(int,float)):
        kpis.append({"key":"ccc","hu":"CCC – Cash Conversion Cycle","en":"CCC – Cash Conversion Cycle","val":ccc,"bmk_hu":"<60 alacsony • >120 magas","bmk_en":"<60 low • >120 high",
                     "status":status_of(ccc,"ccc"), "method_hu":"DSO + DIO − DPO", "method_en":"DSO + DIO − DPO"})

    # section builders
    
    def build_section(doc: Document, lang_code: str, wcn_val=None, nwc_val=None):
            # base styles
            st = doc.styles["Normal"]; st.font.name = "Calibri"; st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri'); st.font.size = Pt(11)

            if lang_code=="hu":
                doc.add_heading("AIRiskMaster (AIRM) – Kockázati riport", 0)
                p = doc.add_paragraph(); p.add_run("Vállalat: ").bold=True; p.add_run(company_name)
                doc.add_paragraph("Időszak: 2024.01.01 – 2024.12.31 (ezer HUF) • Ágazat: " + str(sector).capitalize())
                # overall
                p = doc.add_paragraph(); p.add_run("Össz‑kockázat: ").bold=True; badge(p.add_run(), color); p.add_run(" " + ({"green":"ALACSONY","yellow":"KÖZEPES","red":"MAGAS"}[color])).bold=True
                if isinstance(score,(int,float)): doc.add_paragraph(f"Pontszám: {int(round(score))}/100 • Tartomány: {range_text(color,'hu')} • (CF-vel súlyozva)")
                # rationale + CCC + WCN + drivers
                r_parts = []
                if isinstance(dso,(int,float)): r_parts.append(f"DSO {fmt_num_hu(dso,'days')} a {int(DSO_MAX)} helyett")
                if isinstance(dpo,(int,float)): r_parts.append(f"DPO {fmt_num_hu(dpo,'days')} a {int(DPO_MIN)} helyett")
                if isinstance(ccc,(int,float)): r_parts.append(f"CCC {fmt_num_hu(ccc,'days')}")
                if r_parts: doc.add_paragraph("Indoklás (rövid): " + ", ".join(r_parts) + ".")
                if isinstance(ccc,(int,float)): doc.add_paragraph(f"CCC: {fmt_num_hu(ccc,'days')}")
                if isinstance(wcn_val,(int,float)): doc.add_paragraph(f"WCN (forgótőkeigény): {int(wcn_val):,} eFt".replace(",", " "))
                if isinstance(nwc_val,(int,float)): doc.add_paragraph(f"Nettó forgótőke (NWC): {int(nwc_val):,} eFt".replace(",", " "))
                doc.add_paragraph("Fő kockázati tényezők: • DSO cél felett • DPO ajánlott alatt • CCC magas")
                # 1) fő pénzügyi adatok
                prev_bs = prev.get('bs', {}) if isinstance(prev, dict) else {}
                prev_pl = prev.get('pl', {}) if isinstance(prev, dict) else {}
                doc.add_heading("1) Fő pénzügyi adatok", level=1)
                tb = doc.add_table(rows=1, cols=3); tb.alignment=WD_TABLE_ALIGNMENT.LEFT; tb.style="Light Grid Accent 1"; tb.alignment=WD_TABLE_ALIGNMENT.LEFT
                hdr = tb.rows[0].cells
                hdr[0].text = "Tétel"; hdr[1].text = "Előző év (eFt)"; hdr[2].text = "Tárgyév (eFt)"
                def _row_fin_hu(name, prev, curr):
                    c = tb.add_row().cells
                    c[0].text = name
                    c[1].text = (f"{int(prev):,}".replace(",", " ") if isinstance(prev,(int,float)) else ("-" if prev in (None,"") else str(prev)))
                    c[2].text = (f"{int(curr):,}".replace(",", " ") if isinstance(curr,(int,float)) else ("-" if curr in (None,"") else str(curr)))
                # PL items
                for k in ["Értékesítés nettó árbevétele","Anyagjellegű ráfordítások","Személyi jellegű ráfordítások","Értékcsökkenési leírás","Egyéb bevételek","Egyéb ráfordítások","Adózott eredmény","Üzemi (üzleti) tevékenység eredménye"]:
                    prev_v = prev_pl.get(k); curr_v = pl.get(k)
                    if prev_v in (None, ""): prev_v = raw.get("pl",{}).get(k,{}).get("previous")
                    if curr_v in (None, ""): curr_v = raw.get("pl",{}).get(k,{}).get("current")
                    _row_fin_hu(k, prev_v, curr_v)
                # BS items
                for k in ["Forgóeszközök","Készletek","Követelések","Pénzeszközök","Szállítók","Rövid lejáratú kötelezettségek","Hosszú lejáratú kötelezettségek","Kötelezettségek összesen","Saját tőke"]:
                    prev_v = prev_bs.get(k); curr_v = bs.get(k)
                    if prev_v in (None, ""): prev_v = raw.get("balance",{}).get(k,{}).get("previous")
                    if curr_v in (None, ""): curr_v = raw.get("balance",{}).get(k,{}).get("current")
                    _row_fin_hu(k, prev_v, curr_v)
                doc.add_heading("2) Likviditás és eladósodottság", level=1)
                doc.add_paragraph(f"Current ratio: {fmt_num_hu(cr,'times')}" if cr is not None else "Current ratio: n.a.")
                doc.add_paragraph(f"Quick ratio: {fmt_num_hu(qr,'times')}" if qr is not None else "Quick ratio: n.a.")
                doc.add_paragraph(f"Debt/Equity: {fmt_num_hu(dte,'times')}" if dte is not None else "Debt/Equity: n.a.")
                nfk = ratios.get('Nettó forgótőke (eFt)'); doc.add_paragraph(f"Nettó forgótőke: {int(nfk):,} eFt".replace(",", " ") if isinstance(nfk,(int,float)) else "Nettó forgótőke: n.a.")
                # merged table
                title = doc.add_paragraph(); title.add_run("3) Összevont banki mutatók (benchmark + jelzőlámpa)").bold=True
                tb = doc.add_table(rows=1, cols=5); tb.alignment=WD_TABLE_ALIGNMENT.LEFT; tb.style="Light Grid Accent 1"
                for j,hdr in enumerate(["Mutató / KPI","Érték","Benchmark","Státusz","Értelmezés"]): tb.cell(0,j).text = hdr
                for k in kpis:
                    row = tb.add_row().cells
                    row[0].text = k["hu"]
                    if k["key"] in ("current","quick","de"):
                        row[1].text = fmt_num_hu(k["val"], "times") if k["val"] is not None else "n.a."
                    elif k["key"]=="ccc":
                        row[1].text = fmt_num_hu(k["val"], "days") if isinstance(k["val"],(int,float)) else "n.a."
                    else:
                        row[1].text = fmt_num_hu(k["val"], "days") if isinstance(k["val"],(int,float)) else "n.a."
                    row[2].text = k.get("bmk_hu", k.get("bmk",""))
                    badge(row[3].paragraphs[0].add_run(), k["status"])
                    row[4].text = interp(k["key"], "hu")
                    m_cells = tb.add_row().cells
                    merged = m_cells[0].merge(m_cells[1]).merge(m_cells[2]).merge(m_cells[3]).merge(m_cells[4])
                    para = merged.paragraphs[0]; r=para.add_run("Számítási módszer: "); r.bold=True; para.add_run(k["method_hu"]).italic=True
                # 4) banki ajánlások
                doc.add_heading("4) Banki ajánlások / lépések", level=1)
                doc.add_paragraph("• Faktoring a DSO csökkentésére • Szállítói tárgyalások a DPO hosszabbítására • Rövid lejáratú forgóeszköshitel keret")
                # 5) mellékletek
                built = build_cf_section(doc, lang_code, bs, (prev.get("bs", {}) if isinstance(prev, dict) else {}), pl)
                if not built:
                    doc.add_heading("5) Mellékletek, megjegyzések", level=1)
                    doc.add_paragraph("—")

            else:
                doc.add_heading("AIRiskMaster (AIRM) – Risk Report", 0)
                p = doc.add_paragraph(); p.add_run("Company: ").bold=True; p.add_run(company_name)
                sector_map = {'kereskedelem':'Trade','gyartas':'Manufacturing','gyártás':'Manufacturing','epitoipar':'Construction','építőipar':'Construction','szolgaltatas':'Services','szolgáltatás':'Services','energia':'Energy','agrar':'Agriculture','agrár':'Agriculture'}
                sector_en = sector_map.get(str(sector).lower(), str(sector).capitalize())
                doc.add_paragraph("Period: 2024-01-01 – 2024-12-31 (th HUF) • Sector: " + sector_en)
                p = doc.add_paragraph(); p.add_run("Overall risk: ").bold=True; badge(p.add_run(), color); p.add_run(" " + ({"green":"LOW","yellow":"MODERATE","red":"HIGH"}[color])).bold=True
                if isinstance(score,(int,float)): doc.add_paragraph(f"Score: {int(round(score))}/100 • Range: {range_text(color,'en')} • (CF-weighted)")
                r_parts = []
                if isinstance(dso,(int,float)): r_parts.append(f"DSO {fmt_num_en(dso,'days')} vs {int(DSO_MAX)}")
                if isinstance(dpo,(int,float)): r_parts.append(f"DPO {fmt_num_en(dpo,'days')} vs {int(DPO_MIN)}")
                if isinstance(ccc,(int,float)): r_parts.append(f"CCC {fmt_num_en(ccc,'days')}")
                if r_parts: doc.add_paragraph("Rationale (short): " + ", ".join(r_parts) + ".")
                if isinstance(ccc,(int,float)): doc.add_paragraph(f"CCC: {fmt_num_en(ccc,'days')}")
                if isinstance(wcn_val,(int,float)): doc.add_paragraph(f"WCN (Working Capital Need): {int(wcn_val):,} th HUF".replace(",", " "))
                if isinstance(nwc_val,(int,float)): doc.add_paragraph(f"Net Working Capital (NWC): {int(nwc_val):,} th HUF".replace(",", " "))
                doc.add_paragraph("Key risk drivers: • DSO above target • DPO below recommended • CCC high")

                prev_bs = prev.get('bs', {}) if isinstance(prev, dict) else {}
                prev_pl = prev.get('pl', {}) if isinstance(prev, dict) else {}
                doc.add_heading("1) Key financials", level=1)
                tb = doc.add_table(rows=1, cols=3); tb.alignment=WD_TABLE_ALIGNMENT.LEFT; tb.style="Light Grid Accent 1"; tb.alignment=WD_TABLE_ALIGNMENT.LEFT; tb.style="Light Grid Accent 1"; tb.alignment=WD_TABLE_ALIGNMENT.LEFT
                hdr = tb.rows[0].cells
                hdr[0].text = "Item"; hdr[1].text = "Previous year (th HUF)"; hdr[2].text = "Current year (th HUF)"
                def _row_fin_en(name, prev, curr):
                    c = tb.add_row().cells
                    c[0].text = name
                    c[1].text = (f"{int(prev):,}" if isinstance(prev,(int,float)) else ("-" if prev in (None,"") else str(prev)))
                    c[2].text = (f"{int(curr):,}" if isinstance(curr,(int,float)) else ("-" if curr in (None,"") else str(curr)))
                names = {
                    "Értékesítés nettó árbevétele":"Net sales revenue",
                    "Anyagjellegű ráfordítások":"Material-type expenses",
                    "Személyi jellegű ráfordítások":"Personnel expenses",
                    "Értékcsökkenési leírás":"Depreciation and amortization",
                    "Egyéb bevételek":"Other income",
                    "Egyéb ráfordítások":"Other expenses",
                    "Adózott eredmény":"Profit after tax",
                    "Üzemi (üzleti) tevékenység eredménye":"Operating profit (loss)",
                    "Forgóeszközök":"Current assets",
                    "Készletek":"Inventory",
                    "Követelések":"Receivables",
                    "Pénzeszközök":"Cash and cash equivalents",
                    "Szállítók":"Payables",
                    "Rövid lejáratú kötelezettségek":"Short-term liabilities",
                    "Hosszú lejáratú kötelezettségek":"Long-term liabilities",
                    "Kötelezettségek összesen":"Total liabilities",
                    "Saját tőke":"Equity",
                }
                # PL items
                keys_pl = ["Értékesítés nettó árbevétele","Anyagjellegű ráfordítások","Személyi jellegű ráfordítások","Értékcsökkenési leírás","Egyéb bevételek","Egyéb ráfordítások","Adózott eredmény","Üzemi (üzleti) tevékenység eredménye"]
                for k in keys_pl:
                    prev_v = raw.get("pl",{}).get(k,{}).get("previous")
                    curr_v = raw.get("pl",{}).get(k,{}).get("current")
                    if prev_v in (None,""): prev_v = (prev or {}).get("pl",{}).get(k)
                    if curr_v in (None,""): curr_v = (pl or {}).get(k)
                    _row_fin_en(names.get(k,k), prev_v, curr_v)
                # BS items
                for k in ["Forgóeszközök","Készletek","Követelések","Pénzeszközök","Szállítók","Rövid lejáratú kötelezettségek","Hosszú lejáratú kötelezettségek","Kötelezettségek összesen","Saját tőke"]:
                    r = raw.get("balance",{}).get(k,{})
                    prev_v = (prev or {}).get("bs",{}).get(k)
                    curr_v = bs.get(k)
                    if prev_v in (None, ""): prev_v = raw.get("balance",{}).get(k,{}).get("previous")
                    if curr_v in (None, ""): curr_v = raw.get("balance",{}).get(k,{}).get("current")
                    _row_fin_en(names[k], prev_v, curr_v)
                doc.add_heading("2) Liquidity & leverage", level=1)
                doc.add_paragraph(f"Current ratio: {fmt_num_en(cr,'times')}" if cr is not None else "Current ratio: n.a.")
                doc.add_paragraph(f"Quick ratio: {fmt_num_en(qr,'times')}" if qr is not None else "Quick ratio: n.a.")
                doc.add_paragraph(f"Debt/Equity: {fmt_num_en(dte,'times')}" if dte is not None else "Debt/Equity: n.a.")
                nfk = ratios.get('Nettó forgótőke (eFt)'); doc.add_paragraph(f"Net working capital: {int(nfk):,} th HUF".replace(",", " ") if isinstance(nfk,(int,float)) else "Net working capital: n.a.")

                title = doc.add_paragraph(); title.add_run("3) Combined bank metrics (benchmark + traffic‑light)").bold=True
                tb = doc.add_table(rows=1, cols=5); tb.alignment=WD_TABLE_ALIGNMENT.LEFT; tb.style="Light Grid Accent 1"
                for j,hdr in enumerate(["Metric / KPI","Value","Benchmark","Status","Interpretation"]): tb.cell(0,j).text = hdr
                for k in kpis:
                    row = tb.add_row().cells
                    row[0].text = k["en"]
                    if k["key"] in ("current","quick","de"):
                        row[1].text = fmt_num_en(k["val"], "times") if k["val"] is not None else "n.a."
                    elif k["key"]=="ccc":
                        row[1].text = fmt_num_en(k["val"], "days") if isinstance(k["val"],(int,float)) else "n.a."
                    else:
                        row[1].text = fmt_num_en(k["val"], "days") if isinstance(k["val"],(int,float)) else "n.a."
                    row[2].text = k.get("bmk_en", k.get("bmk",""))
                    badge(row[3].paragraphs[0].add_run(), k["status"])
                    row[4].text = interp(k["key"], "en")
                    m_cells = tb.add_row().cells
                    merged = m_cells[0].merge(m_cells[1]).merge(m_cells[2]).merge(m_cells[3]).merge(m_cells[4])
                    para = merged.paragraphs[0]; r=para.add_run("Calculation method: "); r.bold=True; para.add_run(k["method_en"]).italic=True

                doc.add_heading("4) Bank recommendations / actions", level=1)
                doc.add_paragraph("• Factoring to reduce DSO • Extend supplier terms to lift DPO • Short‑term working‑capital line")
                built = build_cf_section(doc, lang_code, bs, (prev.get("bs", {}) if isinstance(prev, dict) else {}), pl)
                if not built:
                    doc.add_heading("5) Appendices / Notes", level=1)
                    doc.add_paragraph("—")
    # ---- assemble document according to language ----
    doc = Document()
    to_build = ["hu","en"] if (lang in ("both","Both","HU+EN","hu+en")) else ([lang] if lang in ("hu","en") else ["hu"])
    first=True
    for L in to_build:
        if not first: doc.add_page_break()
        build_section(doc, L, wcn, nwc)
        first=False

    out_path = _Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

def _rating_color(val, metric, sector_cfg):
    t = sector_cfg["targets"]
    g = sector_cfg["good"]
    if val is None:
        return ("N/A", "808080", "")
    if metric == "Current ratio":
        target=f">= {t['current_ratio']}"; ok = val>=t['current_ratio']; strong = val>=g['current_ratio_min']
        return ("OK" if ok else "FIGYELEM", "00AA00" if strong else ("55AA55" if ok else "CC0000"), target)
    if metric == "Quick ratio":
        target=f">= {t['quick_ratio']}"; ok = val>=t['quick_ratio']; strong = val>=g['quick_ratio_min']
        return ("OK" if ok else "FIGYELEM", "00AA00" if strong else ("55AA55" if ok else "CC0000"), target)
    if metric == "Debt/Equity":
        target=f"<= {t['debt_to_equity']}"; ok = val<=t['debt_to_equity']; strong = val<=g['debt_to_equity_max']
        return ("OK" if ok else "FIGYELEM", "00AA00" if strong else ("55AA55" if ok else "CC0000"), target)
    if metric == "Vevőállomány forgási ideje (nap)":
        target=f"<= {t['receivables_days']} nap"; ok = val<=t['receivables_days']; strong = val<=g['receivables_days_max']
        return ("OK" if ok else "FIGYELEM", "00AA00" if strong else ("55AA55" if ok else "CC0000"), target)
    if metric == "Készlet forgási ideje (nap)":
        target=f"<= {t['inventory_days']} nap"; ok = val<=t['inventory_days']; strong = val<=g['inventory_days_max']
        return ("OK" if ok else "FIGYELEM", "00AA00" if strong else ("55AA55" if ok else "CC0000"), target)
    if metric == "Szállítói napok (DPO)":
        target=f"<= 75 nap"; ok = val<=75; strong = val<=60
        return ("OK" if ok else "FIGYELEM", "00AA00" if strong else ("55AA55" if ok else "CC0000"), target)
    return ("OK", "000000", "")

def process_file(pdf_path: Path, out_dir: Path, overrides=None, sector='default', lang='hu'):
    text = read_pdf_text(pdf_path)
    bs, pl, raw = parse_financials_with_raw(text)
    # Build previous-year dicts from raw
    prev_bs = {k: raw.get('balance',{}).get(k,{}).get('previous') for k,_ in KEYS_BS}
    prev_pl = {k: (raw.get("pl",{}).get(k,{}).get("previous")) for k,_ in KEYS_PL}
    if overrides:
        for k,v in overrides.get("bs", {}).items():
            if v not in ("", None):
                try: bs[k] = int(v)
                except: pass
        for k,v in overrides.get("pl", {}).items():
            if v not in ("", None):
                try: pl[k] = int(v)
                except: pass
        for k,v in overrides.get("bs_prev", {}).items():
            if v not in ("", None):
                try: prev_bs[k] = int(v)
                except: pass
        for k,v in overrides.get("pl_prev", {}).items():
            if v not in ("", None):
                try: prev_pl[k] = int(v)
                except: pass
    ratios = compute_ratios(bs, pl)
    company_name = pdf_path.stem
    out_dir.mkdir(parents=True, exist_ok=True)
    out_docx = out_dir / f"AIRM_{pdf_path.stem}_riport.docx"
    make_docx(company_name, bs, pl, ratios, out_docx, sector=sector, lang=lang, prev={'bs': prev_bs, 'pl': prev_pl}, raw=raw)
    return {"company": company_name, "bs": bs, "pl": pl, "ratios": ratios, "raw": raw, "docx": str(out_docx)}

def cli():
    import argparse
    ap = argparse.ArgumentParser(description="AIRiskMaster (AIRM) v6.2 – EB PDF -> DOCX (stabil kinyerés + banki mutatók/benchmark)")
    ap.add_argument("pdfs", nargs="+", help="Input EB PDF(ek)")
    ap.add_argument("--out", default="reports", help="Kimeneti mappa")
    ap.add_argument("--overrides", help="JSON fájl a kézi felülírásokhoz")
    ap.add_argument("--sector", default="default", choices=['default','kereskedelem','gyartas','szolgaltatas'], help="Ágazat")
    args = ap.parse_args()
    overrides_map = {}
    if args.overrides:
        with open(args.overrides, "r", encoding="utf-8") as f:
            overrides_map = json.load(f)
    out_dir = Path(args.out)
    results = []
    for p in args.pdfs:
        ov = overrides_map.get(Path(p).name) if overrides_map else None
        res = process_file(Path(p), out_dir, overrides=ov, sector=args.sector)
        results.append(res)
    print(json.dumps(results, ensure_ascii=False, indent=2))

if __name__ == "__main__":
    if len(sys.argv) > 1 and sys.argv[1] != "--gui":
        cli()
    else:
        import tkinter as tk
        from tkinter import filedialog, messagebox
        from tkinter import ttk
        root = tk.Tk()
        root.title("AIRiskMaster (AIRM) v6.2 – EB PDF riport (előnézet + felülírás + benchmark)")
        root.geometry("820x600")
        files = []
        chk_preview_var = tk.BooleanVar(value=True)
        sector_var = tk.StringVar(value="default")
        lang_var = tk.StringVar(value="hu")
        def pick_files():
            paths = filedialog.askopenfilenames(title="Válaszd ki a PDF-eket", filetypes=[("PDF","*.pdf")])
            if paths:
                files.clear(); files.extend(list(paths))
                lbl_files.config(text=f"{len(files)} fájl kiválasztva")
        
        def open_preview_for_file(pdf_path):
            try:
                text = read_pdf_text(Path(pdf_path))
                bs, pl, raw = parse_financials_with_raw(text)
            except Exception as e:
                messagebox.showerror("AIRM", f"Olvasási hiba: {e}"); return None
            top = tk.Toplevel(root); top.title(f"Előnézet és felülírás – {Path(pdf_path).name}"); top.geometry("880x700")
            frm = tk.Frame(top, padx=8, pady=8); frm.pack(fill="both", expand=True)
            tk.Label(frm, text="Mérleg (BS) – előző év (bal) és tárgyév (jobb) – szerkeszthető • + PDF-sor (részlet)").grid(row=0, column=0, columnspan=4, sticky="w")
            entries_bs = {}; entries_bs_prev = {}; r = 1
            for k, _ in KEYS_BS:
                tk.Label(frm, text=k, width=32, anchor="w").grid(row=r, column=0, sticky="w")
                eprev = tk.Entry(frm, width=16); pval = raw["balance"].get(k,{}).get("previous"); eprev.insert(0, "" if pval in (None,"") else str(pval)); eprev.grid(row=r, column=1, sticky="w")
                ecurr = tk.Entry(frm, width=16); val = "" if bs.get(k) is None else str(bs[k]); ecurr.insert(0, val); ecurr.grid(row=r, column=2, sticky="w")
                preview = raw["balance"].get(k,{}).get("line","")
                if not isinstance(preview, str): preview = ""
                tk.Label(frm, text=(preview[:60]+"..." if preview and len(preview)>60 else (preview or "")), fg="#555", anchor="w").grid(row=r, column=3, sticky="w")
                entries_bs_prev[k] = eprev; entries_bs[k] = ecurr; r += 1

            # --- Suppliers line (101) dual field ---
            tk.Label(frm, text="Szállítók (101. sor)", width=32, anchor="w").grid(row=r, column=0, sticky="w")
            pval_sup = raw["balance"].get("Szállítók",{}).get("previous")
            e_sup_prev = tk.Entry(frm, width=16); e_sup_prev.insert(0, "" if pval_sup in (None,"") else str(pval_sup)); e_sup_prev.grid(row=r, column=1, sticky="w")
            val_sup = "" if bs.get("Szállítók") is None else str(bs["Szállítók"])
            e_sup = tk.Entry(frm, width=16); e_sup.insert(0, val_sup); e_sup.grid(row=r, column=2, sticky="w")
            prev_text = raw['balance'].get('Szállítók',{}).get('line','')
            if not isinstance(prev_text, str) or 'None' in str(prev_text): prev_text = ''
            tk.Label(frm, text=(prev_text[:60]+'...' if prev_text and len(prev_text)>60 else (prev_text or '')), fg='#555', anchor='w').grid(row=r, column=3, sticky='w')
            entries_bs_prev["Szállítók"] = e_sup_prev; entries_bs["Szállítók"] = e_sup; r += 1

            tk.Label(frm, text="Eredménykimutatás (PL) – előző év (bal) és tárgyév (jobb) – szerkeszthető • + PDF-sor (részlet)").grid(row=r, column=0, columnspan=4, sticky="w")
            entries_pl = {}; entries_pl_prev = {}; r += 1
            for k, _ in KEYS_PL:
                tk.Label(frm, text=k, width=32, anchor="w").grid(row=r, column=0, sticky="w")
                eprev = tk.Entry(frm, width=16); pval = raw["pl"].get(k,{}).get("previous"); eprev.insert(0, "" if pval in (None,"") else str(pval)); eprev.grid(row=r, column=1, sticky="w")
                ecurr = tk.Entry(frm, width=16); val = "" if pl.get(k) is None else str(pl[k]); ecurr.insert(0, val); ecurr.grid(row=r, column=2, sticky="w")
                preview = raw["pl"].get(k,{}).get("line","")
                tk.Label(frm, text=(preview[:60]+"..." if preview and len(preview)>60 else (preview or "")), fg="#555", anchor="w").grid(row=r, column=3, sticky="w")
                entries_pl_prev[k] = eprev; entries_pl[k] = ecurr; r += 1

            def save_and_close():
                ov = {"bs": {}, "pl": {}, "bs_prev": {}, "pl_prev": {}}
                for k,e in entries_bs.items(): ov["bs"][k] = e.get()
                for k,e in entries_bs_prev.items(): ov["bs_prev"][k] = e.get()
                for k,e in entries_pl.items(): ov["pl"][k] = e.get()
                for k,e in entries_pl_prev.items(): ov["pl_prev"][k] = e.get()
                top.destroy(); 
                result = {"overrides": ov}
                # stash in a closure variable so caller can access
                globals()["_last_preview_result"] = result
                return result

            tk.Button(frm, text="Mentés és riport generálása", command=save_and_close).grid(row=r, column=0, pady=10, sticky="w")
            top.grab_set(); top.wait_window(); 
            return globals().get("_last_preview_result", {}).get("overrides")
        def run_process():
            if not files:
                messagebox.showwarning("AIRM", "Válassz PDF fájlokat!"); return
            out_dir = Path("reports"); out_dir.mkdir(exist_ok=True); log.delete(1.0, tk.END); ok = 0
            for p in files:
                overrides = None
                if chk_preview_var.get():
                    overrides = open_preview_for_file(p)
                    if overrides is None: log.insert(tk.END, f"Megszakítva: {p}\n"); continue
                try:
                    res = process_file(Path(p), out_dir, overrides=overrides, sector=sector_var.get(), lang=lang_var.get())
                    ok += 1; log.insert(tk.END, f"OK: {res['company']} -> {res['docx']}\n")
                except Exception as e:
                    log.insert(tk.END, f"HIBA: {p} -> {e}\n")
            messagebox.showinfo("AIRM", f"Kész! Sikeres riportok: {ok}/{len(files)}")
        frm_top = tk.Frame(root, padx=12, pady=8); frm_top.pack(fill="x")
        tk.Button(frm_top, text="PDF-ek kiválasztása", command=pick_files).pack(side="left")
        lbl_files = tk.Label(frm_top, text="0 fájl kiválasztva"); lbl_files.pack(side="left", padx=10)
        frm_opts = tk.Frame(root, padx=12, pady=4); frm_opts.pack(fill="x")
        tk.Checkbutton(frm_opts, text="Előnézet és kézi felülírás minden fájlra", variable=chk_preview_var).pack(side="left")
        tk.Label(frm_opts, text="  Ágazat:").pack(side="left")
        ttk.Combobox(frm_opts, textvariable=sector_var, values=["default","epitoipar","kereskedelem","gyartas","szolgaltatas","logisztika","elelmiszeripar","kiskereskedelem","nagykereskedelem","it","egeszsegugy","pharma","autoipar","gepipar","elektronika","ingatlan","mezogazdasag"], width=16, state="readonly").pack(side="left", padx=(6,12))
        tk.Label(frm_opts, text="  Nyelv:").pack(side="left")
        ttk.Combobox(frm_opts, textvariable=lang_var, values=["hu","en","both"], width=10, state="readonly").pack(side="left", padx=(6,12))
        frm_actions = tk.Frame(root, padx=12, pady=4); frm_actions.pack(fill="x")
        tk.Button(frm_actions, text="Riportok generálása", command=run_process).pack(side="left")
        log = tk.Text(root, height=16); log.pack(fill="both", expand=True, padx=12, pady=8)
        root.mainloop()