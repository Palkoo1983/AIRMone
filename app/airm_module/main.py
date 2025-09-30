
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pathlib import Path
import time, shutil, sys, importlib.util, traceback, re, json
from typing import Dict, Any, List
import pdfplumber
from docx import Document

BASE_DIR = Path(__file__).resolve().parent
UPLOADS_DIR = BASE_DIR / "uploads"
REPORTS_DIR = BASE_DIR / "reports"
STATIC_DIR = BASE_DIR / "static"
AIRM_DIR = BASE_DIR / "airm_src"

app = FastAPI(title="AIRM Web – golden", version="2025.09.29")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

def ensure_dirs():
    UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)

def import_airm_main():
    if str(AIRM_DIR) not in sys.path:
        sys.path.insert(0, str(AIRM_DIR))
    airm_main_path = AIRM_DIR / "main.py"
    if not airm_main_path.exists():
        raise HTTPException(status_code=500, detail="AIRM main.py nem található az airm_src mappában.")
    try:
        spec = importlib.util.spec_from_file_location("airm_main_module", str(airm_main_path))
        mod = importlib.util.module_from_spec(spec)
        assert spec and spec.loader
        spec.loader.exec_module(mod)  # type: ignore
        return mod
    except Exception as e:
        tb = traceback.format_exc()
        raise HTTPException(status_code=500, detail=f"AIRM import hiba: {e}\n{tb}")

_num_pattern = re.compile(r"-?\d{1,3}(?:[ .]\d{3})*(?:[.,]\d+)?")
def parse_first_number(s: str):
    m = _num_pattern.search(s or "")
    if not m:
        return None
    raw = m.group(0)
    neg = raw.startswith("-")
    raw = raw.replace(" ", "").replace(".", "")
    if "," in raw: raw = raw.replace(",", ".")
    try:
        val = float(raw)
        if neg: val = -abs(val)
        return val
    except Exception:
        return None

def all_docx_text(docx_path: Path) -> str:
    try:
        doc = Document(str(docx_path))
        parts = []
        for p in doc.paragraphs: parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells: parts.append(cell.text)
        for sect in doc.sections:
            try:
                for p in sect.header.paragraphs: parts.append(p.text)
            except: pass
            try:
                for p in sect.footer.paragraphs: parts.append(p.text)
            except: pass
        return "\n".join(parts)
    except Exception:
        return ""

def _find_score_fallback_any_100(text: str):
    t = (text or "").replace("\xa0"," ").lower()
    windows = [t]
    pat = re.compile(r"(\d{1,3})\s*/\s*100")
    for w in windows:
        m = pat.search(w)
        if m:
            try:
                v = float(m.group(1))
                if 0 <= v <= 100: return v
            except: pass
    return None

def find_score(text: str):
    t = (text or "").replace("\xa0"," ").lower()
    i = t.find("pontszám")
    if i != -1:
        v = parse_first_number(t[i:i+80])
        if v is not None: return float(v)
    return None

def find_equity_from_text_or_res(text: str, res_bs: Dict[str,Any]) -> float | None:
    for label in ("Saját tőke","Sajat toke","Equity","own equity"):
        if label in res_bs and res_bs[label] is not None:
            try: return float(res_bs[label])
            except Exception:
                v = parse_first_number(str(res_bs[label]))
                if v is not None: return float(v)
    # fallback to text scan
    t = (text or "").replace("\xa0"," ").lower()
    for key in ["saját tőke","sajat toke","equity","own equity"]:
        i = t.find(key)
        if i != -1:
            v = parse_first_number(t[i:i+120])
            if v is not None: return float(v)
    return None

def _decision_code(decision: str|None) -> str:
    d = (decision or '').strip().lower()
    if 'nem hitelezhető' in d or 'nem hitelezh' in d or d=='reject': return 'REJECT'
    if 'hitelezhető' in d or 'hitelezh' in d or d=='approve': return 'APPROVE'
    return 'UNKNOWN'

def decide_from_metrics(eq, risk):
    if eq is not None and eq < 0: return "nem hitelezhető"
    if risk is not None and risk > 90: return "nem hitelezhető"
    if risk is not None and risk <= 90: return "hitelezhető"
    return "ismeretlen"

def key_labels_from_keys_and_values(keys_seq, values: Dict[str,Any]) -> List[str]:
    labels = []
    if isinstance(keys_seq, (list, tuple)):
        for item in keys_seq:
            if isinstance(item, (list, tuple)) and len(item) >= 1:
                labels.append(str(item[0]))
            else:
                s = str(item); labels.append(s.split(",")[0].strip() or s)
    for k in values.keys():
        if k not in labels: labels.append(k)
    return labels

def sanitize_reports_dir():
    for p in REPORTS_DIR.glob("*.pdf"):
        dest = UPLOADS_DIR / p.name
        try:
            if dest.exists(): dest.unlink()
            shutil.move(str(p), str(dest))
        except Exception: pass

def pdf_to_text(path: Path) -> str:
    try:
        with pdfplumber.open(str(path)) as pdf:
            return "\n".join([(pg.extract_text() or "") for pg in pdf.pages])
    except Exception:
        return ""

@app.get("/", response_class=HTMLResponse)
def root():
    return RedirectResponse(url="/static/index.html")

@app.post("/preview")
async def preview_pdf(file: UploadFile = File(...), sector: str = Form(default="default"), lang: str = Form(default="hu")):
    ensure_dirs()
    if file.content_type not in ("application/pdf", "application/octet-stream"):
        raise HTTPException(status_code=400, detail="Kérlek e-beszámoló PDF-et tölts fel.")
    orig_name = Path(file.filename).name
    stem = "".join(ch for ch in Path(orig_name).stem if ch.isalnum() or ch in ("-","_")).strip() or "file"
    ts = int(time.time())
    saved_name = f"{stem}_{ts}.pdf"
    saved_path = UPLOADS_DIR / saved_name
    with saved_path.open("wb") as out:
        shutil.copyfileobj(file.file, out)

    mod = import_airm_main()
    text = pdf_to_text(saved_path)
    if not text:
        raise HTTPException(status_code=400, detail="Nem sikerült szöveget kinyerni a PDF-ből.")

    try:
        bs_cur, pl_cur, raw = mod.parse_financials_with_raw(text)
    except Exception as e:
        tb = traceback.format_exc()
        raise HTTPException(status_code=500, detail=f"AIRM parser hiba: {e}\n{tb}")

    bs_prev = {}
    for k, info in (raw.get("balance") or {}).items():
        if isinstance(info, dict): bs_prev[k] = info.get("previous")
    pl_prev = {}
    for k, info in (raw.get("pl") or {}).items():
        if isinstance(info, dict): pl_prev[k] = info.get("previous")

    keys_bs = getattr(mod, "KEYS_BS", [])
    keys_pl = getattr(mod, "KEYS_PL", [])
    bs_labels = key_labels_from_keys_and_values(keys_bs, {**bs_cur, **bs_prev})
    pl_labels = key_labels_from_keys_and_values(keys_pl, {**pl_cur, **pl_prev})
    if ("Szállítók" in bs_cur or "Szállítók" in bs_prev) and "Szállítók" not in bs_labels:
        bs_labels.append("Szállítók")

    return JSONResponse({
        "ok": True,
        "saved_pdf": saved_name,
        "sector": sector,
        "lang": lang,
        "bs_labels": bs_labels,
        "pl_labels": pl_labels,
        "bs": bs_cur, "bs_prev": bs_prev,
        "pl": pl_cur, "pl_prev": pl_prev
    })

def _coerce_num(val: Any) -> Any:
    if isinstance(val, (int, float)): return val
    if isinstance(val, str):
        v = val.strip()
        if not v: return v
        v2 = v.replace(" ", "").replace(".", "").replace(",", "")
        if v2.lstrip("-").isdigit():
            try: return int(v2)
            except: return v
        return v
    return val

@app.post("/recalc")
async def recalc(saved_pdf: str = Form(...), sector: str = Form(default="default"), lang: str = Form(default="hu"), overrides_json: str = Form(default="{}")):
    ensure_dirs()
    saved_path = UPLOADS_DIR / saved_pdf
    if not saved_path.exists():
        raise HTTPException(status_code=404, detail="Előnézet fájl nem található (saved_pdf).")
    try:
        overrides: Dict[str, Any] = json.loads(overrides_json or "{}")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Hibás overrides JSON: {e}")
    clean = {}
    for sec in ("bs","bs_prev","pl","pl_prev"):
        sec_dict = overrides.get(sec, {})
        if isinstance(sec_dict, dict):
            filtered = { k: _coerce_num(v) for k,v in sec_dict.items() if str(v).strip() != "" }
            if filtered: clean[sec] = filtered

    mod = import_airm_main()
    try:
        res = mod.process_file(saved_path, REPORTS_DIR, overrides=clean if clean else None, sector=sector, lang=lang)
    except Exception as e:
        tb = traceback.format_exc()
        raise HTTPException(status_code=500, detail=f"AIRM riport-generálás hiba: {e}\n{tb}")

    sanitize_reports_dir()
    out_docx = Path(res.get("docx",""))
    if not out_docx.exists():
        raise HTTPException(status_code=500, detail="AIRM nem hozott létre DOCX kimenetet.")
    text = all_docx_text(out_docx)
    risk = find_score(text) or _find_score_fallback_any_100(text)
    bs2 = res.get("bs", {}) or {}
    eq = find_equity_from_text_or_res(text, bs2)
    decision = decide_from_metrics(eq, risk)
    return JSONResponse({
        "ok": True,
        "decision": decision,
        "decision_code": _decision_code(decision),
        "risk_score": risk,
        "equity_value": eq,
        "docx_file": out_docx.name
    })
